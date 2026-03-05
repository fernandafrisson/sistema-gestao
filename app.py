import streamlit as st
import pandas as pd
import firebase_admin
from firebase_admin import credentials, db
from datetime import datetime, date, timedelta
import plotly.express as px
from geopy.geocoders import Nominatim
import time
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dateutil.relativedelta import relativedelta
import locale
from collections import Counter, defaultdict
import geopandas as gpd
from streamlit_calendar import calendar
import pydeck as pdk

# --- INTERFACE PRINCIPAL ---
st.set_page_config(layout="wide", page_title="Sistema de Gestao - CCZ", page_icon="🏥")

# --- CSS GLOBAL DO SISTEMA ---
st.markdown("""
<style>
    /* === RESET & BASE === */
    @import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&display=swap');

    html, body, [class*="css"] {
        font-family: 'DM Sans', sans-serif;
    }
    
    /* === SIDEBAR === */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #0F2940 0%, #1B4F72 100%);
    }
    [data-testid="stSidebar"] * {
        color: #D6EAF8 !important;
    }
    [data-testid="stSidebar"] .stButton > button {
        background: rgba(255,255,255,0.1) !important;
        border: 1px solid rgba(255,255,255,0.2) !important;
        color: white !important;
        border-radius: 10px !important;
        transition: all 0.2s ease !important;
    }
    [data-testid="stSidebar"] .stButton > button:hover {
        background: rgba(255,255,255,0.2) !important;
        border-color: rgba(255,255,255,0.4) !important;
    }

    /* === TABS GLOBAIS === */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background: #F0F3F8;
        padding: 6px;
        border-radius: 14px;
    }
    .stTabs [data-baseweb="tab"] {
        border-radius: 10px;
        padding: 10px 20px;
        background: transparent;
        font-weight: 500;
        font-size: 0.9rem;
        color: #566573;
    }
    .stTabs [aria-selected="true"] {
        background: #1B4F72 !important;
        color: white !important;
        font-weight: 600;
        box-shadow: 0 2px 8px rgba(27,79,114,0.3);
    }

    /* === CARDS === */
    .sys-card {
        background: #ffffff;
        border: 1px solid #E5E8EB;
        border-radius: 14px;
        padding: 24px 28px;
        margin-bottom: 20px;
        box-shadow: 0 2px 12px rgba(0,0,0,0.04);
        transition: box-shadow 0.2s ease;
    }
    .sys-card:hover {
        box-shadow: 0 4px 20px rgba(0,0,0,0.08);
    }
    .sys-card-title {
        font-size: 1.1rem;
        font-weight: 600;
        color: #1B4F72;
        margin-bottom: 16px;
        padding-bottom: 10px;
        border-bottom: 2px solid #D4E6F1;
        display: flex;
        align-items: center;
        gap: 8px;
    }

    /* === HEADERS DE MODULO === */
    .mod-header {
        background: linear-gradient(135deg, #1B4F72 0%, #2E86C1 100%);
        border-radius: 16px;
        padding: 28px 32px;
        margin-bottom: 28px;
        color: white;
    }
    .mod-header h2 {
        margin: 0 0 6px 0;
        font-size: 1.6rem;
        font-weight: 700;
        color: white !important;
        letter-spacing: -0.3px;
    }
    .mod-header p {
        margin: 0;
        font-size: 0.92rem;
        opacity: 0.85;
        color: #D6EAF8;
    }

    /* === METRICAS === */
    .metric-row {
        display: flex;
        gap: 16px;
        margin-bottom: 24px;
    }
    .metric-box {
        flex: 1;
        background: #F8F9FA;
        border: 1px solid #E5E8EB;
        border-radius: 12px;
        padding: 18px 20px;
        text-align: center;
    }
    .metric-number {
        font-size: 1.9rem;
        font-weight: 700;
        color: #1B4F72;
        line-height: 1;
    }
    .metric-label {
        font-size: 0.78rem;
        color: #85929E;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        margin-top: 4px;
    }

    /* === BADGES === */
    .sys-badge {
        display: inline-block;
        padding: 4px 14px;
        border-radius: 20px;
        font-size: 0.78rem;
        font-weight: 600;
        letter-spacing: 0.5px;
        text-transform: uppercase;
    }
    .badge-green { background: #D5F5E3; color: #1E8449; }
    .badge-blue { background: #D4E6F1; color: #1B4F72; }
    .badge-orange { background: #FDEBD0; color: #B9770E; }
    .badge-red { background: #FADBD8; color: #C0392B; }
    .badge-purple { background: #E8DAEF; color: #6C3483; }
    .badge-gray { background: #EAECEE; color: #566573; }

    /* === INFO GRID === */
    .info-grid {
        display: grid;
        grid-template-columns: 1fr 1fr;
        gap: 10px 24px;
        margin-top: 12px;
    }
    .info-item { padding: 8px 0; }
    .info-label {
        font-size: 0.75rem;
        font-weight: 600;
        color: #85929E;
        text-transform: uppercase;
        letter-spacing: 0.6px;
        margin-bottom: 2px;
    }
    .info-value {
        font-size: 0.95rem;
        color: #2C3E50;
        font-weight: 500;
    }

    /* === TAGS/CHIPS === */
    .sys-tag {
        display: inline-block;
        background: #EBF5FB;
        color: #2471A3;
        padding: 5px 12px;
        border-radius: 8px;
        font-size: 0.85rem;
        margin: 3px 4px 3px 0;
        font-weight: 500;
    }
    .sys-chip {
        display: inline-block;
        background: #EAF2F8;
        border: 1px solid #AED6F1;
        color: #1A5276;
        padding: 5px 14px;
        border-radius: 20px;
        font-size: 0.83rem;
        margin: 4px 6px 4px 0;
        font-weight: 500;
    }

    /* === DIVIDER === */
    .sys-divider {
        height: 1px;
        background: linear-gradient(90deg, transparent, #D4E6F1, transparent);
        margin: 16px 0;
        border: none;
    }

    /* === TIMELINE/HIST CARDS === */
    .hist-card {
        background: #FAFBFC;
        border: 1px solid #E5E8EB;
        border-left: 4px solid #2E86C1;
        border-radius: 10px;
        padding: 20px 24px;
        margin-bottom: 14px;
    }

    /* === EMPTY STATE === */
    .empty-state {
        text-align: center;
        padding: 48px 24px;
        color: #85929E;
    }
    .empty-state .icon {
        font-size: 3rem;
        margin-bottom: 12px;
        opacity: 0.5;
    }
    .empty-state p {
        font-size: 1rem;
        margin: 0;
    }

    /* === LOGIN === */
    .login-container {
        max-width: 420px;
        margin: 60px auto;
        text-align: center;
    }
    .login-logo {
        font-size: 3.5rem;
        margin-bottom: 8px;
    }
    .login-title {
        font-size: 1.8rem;
        font-weight: 700;
        color: #1B4F72;
        margin-bottom: 4px;
    }
    .login-subtitle {
        font-size: 0.95rem;
        color: #85929E;
        margin-bottom: 32px;
    }

    /* === PAINEL CARDS DE MODULO === */
    .module-card {
        background: #ffffff;
        border: 1px solid #E5E8EB;
        border-radius: 16px;
        padding: 28px 24px;
        text-align: center;
        transition: all 0.25s ease;
        cursor: pointer;
        box-shadow: 0 2px 8px rgba(0,0,0,0.04);
    }
    .module-card:hover {
        box-shadow: 0 8px 24px rgba(0,0,0,0.1);
        transform: translateY(-2px);
        border-color: #2E86C1;
    }
    .module-card .icon { font-size: 2.4rem; margin-bottom: 10px; }
    .module-card .title { font-size: 1rem; font-weight: 600; color: #2C3E50; }
    .module-card .desc { font-size: 0.82rem; color: #85929E; margin-top: 4px; }

    /* === FICHA FUNCIONARIO === */
    .ficha-card {
        background: #ffffff;
        border: 1px solid #E5E8EB;
        border-radius: 14px;
        padding: 24px;
        text-align: center;
    }
    .ficha-avatar {
        width: 80px;
        height: 80px;
        border-radius: 50%;
        background: linear-gradient(135deg, #1B4F72, #2E86C1);
        color: white;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 2rem;
        font-weight: 700;
        margin: 0 auto 14px;
    }
    .ficha-name {
        font-size: 1.15rem;
        font-weight: 600;
        color: #2C3E50;
        margin-bottom: 4px;
    }
    .ficha-role {
        font-size: 0.88rem;
        color: #85929E;
        margin-bottom: 16px;
    }
    .ficha-detail {
        display: flex;
        justify-content: space-between;
        padding: 8px 0;
        border-bottom: 1px solid #F0F3F8;
        font-size: 0.88rem;
    }
    .ficha-detail-label {
        color: #85929E;
        font-weight: 500;
    }
    .ficha-detail-value {
        color: #2C3E50;
        font-weight: 600;
    }

    /* === NOTES/OBS BOX === */
    .obs-box {
        margin-top: 12px;
        padding: 10px 14px;
        background: #FEF9E7;
        border-radius: 8px;
        font-size: 0.88rem;
        color: #7D6608;
    }

    /* === BOTOES GLOBAIS === */
    .stButton > button[kind="primary"] {
        border-radius: 10px !important;
        font-weight: 600 !important;
    }
    .stButton > button {
        border-radius: 10px !important;
    }

    /* === DATAFRAME === */
    .stDataFrame {
        border-radius: 12px;
        overflow: hidden;
    }
</style>
""", unsafe_allow_html=True)

# --- USUÁRIOS PARA LOGIN (Exemplo) ---
USERS = {
    "admin": "admin123",
    "taylan": "taylan123",
    "fernanda": "fernanda123"
}

# --- CONFIGURAÇÃO DO FIREBASE ---
try:
    if not firebase_admin._apps:
        if 'firebase_credentials' in st.secrets:
            cred_dict = dict(st.secrets["firebase_credentials"])
            cred = credentials.Certificate(cred_dict)
        else:
            cred = credentials.Certificate("denuncias-48660-firebase-adminsdk-fbsvc-9f27fef1c8.json")

        firebase_admin.initialize_app(cred, {
            'databaseURL': 'https://denuncias-48660-default-rtdb.firebaseio.com/'
        })
except Exception as e:
    st.error(f"Erro ao inicializar o Firebase: {e}. Verifique as suas credenciais.")


# --- FUNÇÕES GLOBAIS DE DADOS E UTILITÁRIAS ---

def formatar_nome(nome_completo):
    """Retorna o primeiro e o segundo nome de um nome completo."""
    if not isinstance(nome_completo, str):
        return ""
    partes = nome_completo.split()
    if len(partes) > 1:
        return f"{partes[0]} {partes[1]}"
    return partes[0] if partes else ""

# ### NOVA FUNÇÃO DE LOG DE ATIVIDADE ###
def log_atividade(usuario, acao, detalhes=""):
    """
    Registra uma ação do usuário no banco de dados.
    """
    try:
        ref = db.reference('logs_de_atividade')
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        log_id = str(int(time.time() * 1000))
        
        ref.child(log_id).set({
            "usuario": usuario,
            "acao": acao,
            "detalhes": detalhes,
            "timestamp": timestamp
        })
    except Exception as e:
        st.error(f"Erro ao registrar log de atividade: {e}")


@st.cache_data
def carregar_dados_firebase(node):
    """Carrega dados de um nó do Firebase e retorna como DataFrame."""
    try:
        ref = db.reference(f'/{node}')
        data = ref.get()
        if data:
            if isinstance(data, dict):
                df = pd.DataFrame.from_dict(data, orient='index')
                if 'id' not in df.columns:
                    df['id'] = df.index
                return df
            elif isinstance(data, list) and all(isinstance(item, dict) for item in data if item):
                return pd.DataFrame([item for item in data if item])
            return pd.DataFrame()
        return pd.DataFrame()
    except Exception as e:
        st.error(f"Erro ao carregar dados do nó '{node}': {e}")
        return pd.DataFrame()

@st.cache_data
def carregar_quarteiroes_csv():
    """Carrega lista de quarteirões de um CSV no GitHub."""
    url_csv = 'https://raw.githubusercontent.com/fernandafrisson/sistema-gestao/main/Quarteirao.csv'
    try:
        df_quarteiroes = pd.read_csv(url_csv, header=None, encoding='latin-1')
        quarteiroes_lista = sorted(df_quarteiroes[0].astype(str).unique().tolist())
        return quarteiroes_lista
    except Exception as e:
        st.error(f"Não foi possível carregar a lista de quarteirões. Erro: {e}")
        return []

@st.cache_data
def carregar_geo_kml():
    """Carrega dados de geolocalização de um arquivo KML no GitHub."""
    url_kml = 'https://raw.githubusercontent.com/fernandafrisson/sistema-gestao/main/Quadras%20de%20Guar%C3%A1.kml'
    try:
        gdf = gpd.read_file(url_kml, driver='KML')
        pontos = []
        for index, row in gdf.iterrows():
            quadra_nome = row['Name']
            if row['geometry'] is not None and hasattr(row['geometry'], 'geom_type'):
                if row['geometry'].geom_type == 'Point':
                    lon, lat = row['geometry'].x, row['geometry'].y
                else:
                    centroid = row['geometry'].centroid
                    lon, lat = centroid.x, centroid.y
                pontos.append({'quadra': str(quadra_nome), 'lat': lat, 'lon': lon})
        df_geo = pd.DataFrame(pontos)
        df_geo.dropna(subset=['lat', 'lon'], inplace=True)
        return df_geo
    except Exception as e:
        st.error(f"Não foi possível carregar os dados de geolocalização do KML. Erro: {e}")
        return pd.DataFrame()

# --- FUNÇÕES DE GERAÇÃO DE RELATÓRIOS .DOCX ---

def create_abonada_word_report(data):
    """Gera um relatório de Falta Abonada em formato .docx."""
    def format_date_pt(dt):
        months = ("Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho", "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro")
        return f"{dt.day} de {months[dt.month - 1]} de {dt.year}"
    document = Document()
    black_color = RGBColor(0, 0, 0)
    def add_black_run(p, text, bold=False, size=11):
        run = p.add_run(text)
        run.bold = bold
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(size)
        font.color.rgb = black_color
        p.paragraph_format.space_after = Pt(0)
    for text in ["Fundo Municipal de Saúde", "Prefeitura Municipal da Estância Turística de Guaratinguetá", "São Paulo", "Secretaria Municipal da Saúde"]:
        p = document.add_paragraph()
        add_black_run(p, text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo = document.add_heading('FALTA ABONADA', level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.color.rgb = black_color
    titulo.paragraph_format.space_before = Pt(18)
    titulo.paragraph_format.space_after = Pt(18)
    p_nome = document.add_paragraph(); add_black_run(p_nome, 'Nome: '); add_black_run(p_nome, data.get('nome', ''), bold=True)
    p_funcao = document.add_paragraph(); add_black_run(p_funcao, 'Função: '); add_black_run(p_funcao, data.get('funcao', ''), bold=True)
    p_unidade = document.add_paragraph(); add_black_run(p_unidade, 'Unidade de Trabalho: '); add_black_run(p_unidade, data.get('unidade', ''), bold=True)
    solicitacao_text = f"Solicito que a minha falta ao serviço seja abonada no dia: {data.get('data_abonada', '')}"
    p_solicitacao = document.add_paragraph(); add_black_run(p_solicitacao, solicitacao_text)
    p_solicitacao.paragraph_format.space_before = Pt(18)
    p_solicitacao.paragraph_format.space_after = Pt(18)
    data_atual_formatada = format_date_pt(date.today())
    p_data = document.add_paragraph(f"Guaratinguetá, {data_atual_formatada}"); p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p_data.runs: run.font.color.rgb = black_color
    p_data.paragraph_format.space_after = Pt(36)
    p_ass1 = document.add_paragraph('____________________________'); p_ass1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_lab1 = document.add_paragraph('Assinatura do Servidor'); p_lab1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ass2 = document.add_paragraph('_____________________________'); p_ass2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_lab2 = document.add_paragraph('Assinatura da Chefia Imediata'); p_lab2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_lab2.paragraph_format.space_after = Pt(18)
    p_info = document.add_paragraph(); add_black_run(p_info, 'Informação da Seção de Pessoal:', bold=True)
    add_black_run(document.add_paragraph(), "Refere-se à:       1ª (   )     2ª (   )   3ª (   ) do Primeiro Semestre de: ____________")
    add_black_run(document.add_paragraph(), "           1ª (   )     2ª (   )   3ª (   ) do Segundo Semestre de: ____________")
    p_visto = document.add_paragraph("      ___________________________________________");
    p_visto_label = document.add_paragraph("          (visto do funcionário da seção de pessoal)")
    p_abone = document.add_paragraph("                Abone-se: _____/_____/______")
    p_abone.paragraph_format.space_after = Pt(18)
    p_secretario_sig = document.add_paragraph("_________________________________"); p_secretario_sig.alignment = 1
    p_secretario_label = document.add_paragraph("Secretário Municipal da Saúde"); p_secretario_label.alignment = 1
    for p in document.paragraphs:
        if not p.runs:
            p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            if run.font.color.rgb is None:
                run.font.color.rgb = black_color
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def create_word_report(data):
    """Gera um relatório de Inspeção Zoossanitária em formato .docx."""
    document = Document()
    style = document.styles['Normal']; font = style.font; font.name = 'Calibri'; font.size = Pt(11)
    titulo = document.add_heading('RELATÓRIO DE INSPEÇÃO ZOOSSANITÁRIA', level=1); titulo.alignment = 1
    try: data_obj = datetime.strptime(data.get('data_denuncia', ''), '%Y-%m-%d'); data_formatada = data_obj.strftime('%d/%m/%Y')
    except (ValueError, TypeError): data_formatada = "Data não informada"
    p_data = document.add_paragraph(data_formatada); p_data.alignment = 2
    document.add_paragraph('Vigilância Epidemiológica')
    p = document.add_paragraph(); p.add_run('Responsável: ').bold = True; p.add_run(str(data.get('responsavel_atendimento', '')))
    endereco_completo = f"{data.get('logradouro', '')}, {data.get('numero', '')} - {data.get('bairro', '')}"
    p = document.add_paragraph(); p.add_run('Endereço: ').bold = True; p.add_run(endereco_completo)
    document.add_paragraph(); p = document.add_paragraph(); p.add_run('Relato da Situação: ').bold = True
    document.add_paragraph(str(data.get('detalhes_denuncia', '')))
    document.add_paragraph(); p = document.add_paragraph(); p.add_run('Situação Encontrada: ').bold = True
    document.add_paragraph(str(data.get('relatorio_atendimento', '')))
    document.add_paragraph(); p = document.add_paragraph(); p.add_run('Conclusão: ').bold = True
    document.add_paragraph(str(data.get('conclusao_atendimento', '')))
    footer = document.sections[0].footer; footer_para = footer.paragraphs[0]
    footer_para.text = ("PREFEITURA MUNICIPAL DA ESTÂNCIA TURÍSTICA DE GUARATINGUETÁ/SP\n"
                        "Secretaria Municipal de Saúde - Fundo Municipal de Saúde\n"
                        "Rua Jacques Felix, 02 – São Gonçalo - Guaratinguetá/SP - CEP 12.502-180\n"
                        "Telefone / Fax: (12) 3123-2900 - e-mail: ccz@guaratingueta.sp.gov.br")
    footer_para.alignment = 1
    font_footer = footer_para.style.font
    font_footer.name = 'Arial'; font_footer.size = Pt(8)
    buffer = io.BytesIO(); document.save(buffer); buffer.seek(0)
    return buffer.getvalue()

def create_boletim_word_report(data_boletim):
    """Gera um relatório do boletim diário em formato .docx."""
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    titulo = document.add_heading('BOLETIM DE PROGRAMAÇÃO DIÁRIA', level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    try:
        locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
    except locale.Error:
        locale.setlocale(locale.LC_TIME, '') 

    data_formatada = pd.to_datetime(data_boletim.get('data')).strftime('%d de %B de %Y')
    p_data = document.add_paragraph(data_formatada.title())
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_data.paragraph_format.space_after = Pt(18)

    document.add_heading('Informações Gerais', level=2)
    p_bairros = document.add_paragraph()
    p_bairros.add_run('Bairros Trabalhados: ').bold = True
    p_bairros.add_run(data_boletim.get('bairros', 'Não informado'))

    p_atividades = document.add_paragraph()
    p_atividades.add_run('Atividades Gerais: ').bold = True
    p_atividades.add_run(', '.join(data_boletim.get('atividades_gerais', ['Nenhuma'])))

    p_motoristas = document.add_paragraph()
    p_motoristas.add_run('Motorista(s): ').bold = True
    motoristas_formatados = [formatar_nome(nome) for nome in data_boletim.get('motoristas', [])]
    p_motoristas.add_run(', '.join(motoristas_formatados) if motoristas_formatados else 'Nenhum')
    
    document.add_heading('Ausências do Dia', level=2)
    faltas_manha = data_boletim.get('faltas_manha', {})
    nomes_manha = [formatar_nome(nome) for nome in faltas_manha.get('nomes', [])]
    p_faltas_m = document.add_paragraph()
    p_faltas_m.add_run('Manhã: ').bold = True
    p_faltas_m.add_run(f"{', '.join(nomes_manha) if nomes_manha else 'Nenhuma'} - Motivo: {faltas_manha.get('motivo', 'N/A')}")

    faltas_tarde = data_boletim.get('faltas_tarde', {})
    nomes_tarde = [formatar_nome(nome) for nome in faltas_tarde.get('nomes', [])]
    p_faltas_t = document.add_paragraph()
    p_faltas_t.add_run('Tarde: ').bold = True
    p_faltas_t.add_run(f"{', '.join(nomes_tarde) if nomes_tarde else 'Nenhuma'} - Motivo: {faltas_tarde.get('motivo', 'N/A')}")
    
    def adicionar_secao_turno(turno_nome, equipes):
        document.add_heading(f'Turno da {turno_nome}', level=2)
        if not equipes or not isinstance(equipes, list):
            document.add_paragraph('Nenhuma equipe registrada para este turno.')
            return
            
        for i, equipe in enumerate(equipes):
            document.add_heading(f'Equipe {i + 1}', level=3)
            
            membros_formatados = [formatar_nome(nome) for nome in equipe.get('membros', [])]
            p_membros = document.add_paragraph()
            p_membros.add_run('Membros: ').bold = True
            p_membros.add_run(', '.join(membros_formatados) if membros_formatados else 'Nenhum')

            p_atividades_eq = document.add_paragraph()
            p_atividades_eq.add_run('Atividades: ').bold = True
            p_atividades_eq.add_run(', '.join(equipe.get('atividades', ['Nenhuma'])))

            p_quarteiroes = document.add_paragraph()
            p_quarteiroes.add_run('Quarteirões: ').bold = True
            p_quarteiroes.add_run(', '.join(equipe.get('quarteiroes', ['Nenhum'])))
            p_quarteiroes.paragraph_format.space_after = Pt(12)

    adicionar_secao_turno("Manhã", data_boletim.get('equipes_manha'))
    adicionar_secao_turno("Tarde", data_boletim.get('equipes_tarde'))
    
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# --- FUNÇÕES ESPECÍFICAS DO MÓDULO RH ---

def calcular_status_ferias_saldo(employee_row, all_folgas_df):
    """Calcula o status de férias de um funcionário."""
    try:
        today = date.today()
        if 'data_admissao' not in employee_row or pd.isna(employee_row['data_admissao']):
            return "Admissão Inválida", "Erro", "ERROR"

        data_admissao = pd.to_datetime(employee_row['data_admissao']).date()
        
        ferias_do_funcionario = pd.DataFrame()
        if not all_folgas_df.empty and 'id_funcionario' in all_folgas_df.columns:
            ferias_do_funcionario = all_folgas_df[(all_folgas_df['id_funcionario'] == str(employee_row['id'])) & (all_folgas_df['tipo'] == 'Férias')].copy()
            if not ferias_do_funcionario.empty:
                ferias_do_funcionario['data_inicio'] = pd.to_datetime(ferias_do_funcionario['data_inicio']).dt.date
                ferias_do_funcionario['data_fim'] = pd.to_datetime(ferias_do_funcionario['data_fim']).dt.date
                
                for _, ferias in ferias_do_funcionario.iterrows():
                    if ferias['data_inicio'] <= today <= ferias['data_fim']:
                        return f"Em gozo desde {ferias['data_inicio'].strftime('%d/%m/%Y')}", "EM FÉRIAS", "ON_VACATION"

        periodos_pendentes = []
        periodo_aquisitivo_inicio = data_admissao
        
        proximo_periodo_aquisitivo = data_admissao 
        
        while True:
            periodo_aquisitivo_fim = periodo_aquisitivo_inicio + relativedelta(years=1) - relativedelta(days=1)
            
            if today < periodo_aquisitivo_fim:
                proximo_periodo_aquisitivo = periodo_aquisitivo_inicio
                break

            periodo_concessivo_fim = periodo_aquisitivo_fim + relativedelta(years=1)
            
            dias_gozados = 0
            if not ferias_do_funcionario.empty:
                ferias_neste_periodo = ferias_do_funcionario[
                    (ferias_do_funcionario['data_inicio'] > periodo_aquisitivo_fim) & 
                    (ferias_do_funcionario['data_inicio'] <= periodo_concessivo_fim)
                ]
                if not ferias_neste_periodo.empty:
                    dias_gozados = sum((fim - inicio).days + 1 for inicio, fim in zip(ferias_neste_periodo['data_inicio'], ferias_neste_periodo['data_fim']))
            
            if dias_gozados < 30:
                periodos_pendentes.append({
                    "inicio_aq": periodo_aquisitivo_inicio,
                    "fim_aq": periodo_aquisitivo_fim,
                    "fim_con": periodo_concessivo_fim,
                    "dias_gozados": dias_gozados
                })
            
            periodo_aquisitivo_inicio += relativedelta(years=1)
            if periodo_aquisitivo_inicio.year > today.year + 2:
                proximo_periodo_aquisitivo = periodo_aquisitivo_inicio
                break

        if len(periodos_pendentes) >= 2:
            periodo_mais_antigo = periodos_pendentes[0]
            fim_concessivo_antigo = periodo_mais_antigo["fim_con"]
            
            if today >= fim_concessivo_antigo:
                return f"Venceu em: {fim_concessivo_antigo.strftime('%d/%m/%Y')}", "RISCO: 2ª FÉRIAS VENCIDA!", "RISK_EXPIRING"
            if (fim_concessivo_antigo - today).days <= 90:
                return f"Vencimento em: {fim_concessivo_antigo.strftime('%d/%m/%Y')}", "RISCO: VENCIMENTO DE 2ª FÉRIAS!", "RISK_EXPIRING"

        if periodos_pendentes:
            periodo_a_reportar = periodos_pendentes[0]
            ref_periodo_str = f"{periodo_a_reportar['inicio_aq'].strftime('%d/%m/%Y')} a {periodo_a_reportar['fim_aq'].strftime('%d/%m/%Y')}"
            
            if periodo_a_reportar['dias_gozados'] > 0:
                return ref_periodo_str, f"Parcialmente Agendada ({periodo_a_reportar['dias_gozados']}/30)", "SCHEDULED"
            else:
                return ref_periodo_str, "PENDENTE DE AGENDAMENTO", "PENDING"
                
        aq_inicio = proximo_periodo_aquisitivo
        aq_fim = aq_inicio + relativedelta(years=1) - relativedelta(days=1)
        if today <= aq_fim:
            return f"{aq_inicio.strftime('%d/%m/%Y')} a {aq_fim.strftime('%d/%m/%Y')}", "Em Aquisição", "ACQUIRING"

        return "N/A", "Em dia", "OK"

    except Exception as e:
        return "Erro de Cálculo", f"Erro: {e}", "ERROR"

def get_abonadas_ano(employee_id, all_folgas_df):
    """Retorna o número de faltas abonadas no ano corrente para um funcionário."""
    try:
        current_year = date.today().year
        if all_folgas_df.empty or 'id_funcionario' not in all_folgas_df.columns:
            return 0
        abonadas_funcionario = all_folgas_df[(all_folgas_df['id_funcionario'] == str(employee_id)) & (all_folgas_df['tipo'] == 'Abonada') & (pd.to_datetime(all_folgas_df['data_inicio']).dt.year == current_year)]
        return len(abonadas_funcionario)
    except Exception:
        return 0

def get_datas_abonadas_ano(employee_id, all_folgas_df):
    """Retorna as datas das faltas abonadas no ano corrente."""
    try:
        current_year = date.today().year
        if all_folgas_df.empty or 'id_funcionario' not in all_folgas_df.columns:
            return []
        
        abonadas_df = all_folgas_df[
            (all_folgas_df['id_funcionario'] == str(employee_id)) & 
            (all_folgas_df['tipo'] == 'Abonada') & 
            (pd.to_datetime(all_folgas_df['data_inicio']).dt.year == current_year)
        ]
        
        if abonadas_df.empty:
            return []
            
        return [pd.to_datetime(d).strftime('%d/%m/%Y') for d in abonadas_df['data_inicio']]
    except Exception:
        return []

def get_ultimas_ferias(employee_id, all_folgas_df):
    """Retorna a data de início do último período de férias registrado."""
    try:
        if all_folgas_df.empty or 'id_funcionario' not in all_folgas_df.columns:
            return "Nenhum registro"
        ferias_do_funcionario = all_folgas_df[(all_folgas_df['id_funcionario'] == str(employee_id)) & (all_folgas_df['tipo'] == 'Férias')].copy()
        if ferias_do_funcionario.empty:
            return "Nenhuma férias registrada"
        ferias_do_funcionario['data_inicio'] = pd.to_datetime(ferias_do_funcionario['data_inicio'])
        ultima_ferias = ferias_do_funcionario.sort_values(by='data_inicio', ascending=False).iloc[0]
        return ultima_ferias['data_inicio'].strftime('%d/%m/%Y')
    except Exception:
        return "Erro"

# --- MÓDULOS DA APLICAÇÃO ---

def modulo_rh():
    """Renderiza a pagina do modulo de Recursos Humanos."""
    st.markdown("""
        <div class="mod-header">
            <h2>👥 Recursos Humanos</h2>
            <p>Gestao de equipe, ferias, abonadas e cadastro de funcionarios</p>
        </div>
    """, unsafe_allow_html=True)
    df_funcionarios = carregar_dados_firebase('funcionarios')
    df_folgas = carregar_dados_firebase('folgas_ferias')

    if not df_funcionarios.empty:
        nome_map = {formatar_nome(nome): nome for nome in df_funcionarios['nome']}
        lista_nomes_curtos = sorted(list(nome_map.keys()))
    else:
        nome_map = {}
        lista_nomes_curtos = []
    
    tab_rh1, tab_rh2, tab_rh3 = st.tabs(["✈️ Férias e Abonadas", "👥 Visualizar Equipe", "👨‍💼 Gerenciar Funcionários"])
    
    with tab_rh1:
        st.subheader("Registro de Férias e Abonadas")
        if lista_nomes_curtos:
            nome_curto_selecionado = st.selectbox("Selecione o Funcionário", lista_nomes_curtos)
            tipo_evento = st.selectbox("Tipo de Evento", ["Férias", "Abonada"], key="tipo_evento_selector")
            
            if 'doc_data' not in st.session_state:
                st.session_state.doc_data = None

            with st.form("folgas_ferias_form", clear_on_submit=True):
                if tipo_evento == "Férias":
                    st.write("Período de Férias:")
                    col1, col2 = st.columns(2)
                    with col1:
                        data_inicio = st.date_input("Data de Início")
                    with col2:
                        data_fim = st.date_input("Data de Fim")
                else:
                    st.write("Data da Abonada:")
                    data_inicio = st.date_input("Data")
                    data_fim = data_inicio
                
                submit_evento = st.form_submit_button("Registrar Evento")
                
                if submit_evento:
                    nome_completo = nome_map[nome_curto_selecionado]
                    if tipo_evento == "Férias" and data_inicio > data_fim:
                        st.error("A data de início não pode ser posterior à data de fim.")
                    else:
                        try:
                            id_funcionario = df_funcionarios[df_funcionarios['nome'] == nome_completo]['id'].iloc[0]
                            evento_id = str(int(time.time() * 1000))
                            ref = db.reference(f'folgas_ferias/{evento_id}')
                            ref.set({'id_funcionario': id_funcionario,'nome_funcionario': nome_completo,'tipo': tipo_evento,'data_inicio': data_inicio.strftime("%Y-%m-%d"),'data_fim': data_fim.strftime("%Y-%m-%d")})
                            
                            log_atividade(st.session_state.get('username'), f"Registrou {tipo_evento}", f"Funcionário: {nome_completo}, Período: {data_inicio} a {data_fim}")

                            st.success(f"{tipo_evento} para {nome_completo} registrado com sucesso!")
                            
                            if tipo_evento == "Abonada":
                                dados_func = df_funcionarios[df_funcionarios['id'] == id_funcionario].iloc[0]
                                doc_data = {'nome': dados_func.get('nome', ''),'funcao': dados_func.get('funcao', ''),'unidade': dados_func.get('unidade_trabalho', ''),'data_abonada': data_inicio.strftime('%d-%m-%Y'),}
                                st.session_state.doc_data = doc_data
                            else:
                                st.session_state.doc_data = None
                            
                            st.cache_data.clear()
                            st.rerun()
                        except Exception as e:
                            st.error(f"Erro ao registrar evento: {e}")

            if st.session_state.doc_data:
                word_bytes = create_abonada_word_report(st.session_state.doc_data)
                st.download_button(label="📥 Baixar Requerimento de Abonada (.docx)",data=word_bytes,file_name=f"Abonada_{st.session_state.doc_data['nome']}_{st.session_state.doc_data['data_abonada']}.docx",mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.info("Nenhum funcionário cadastrado.")
        st.divider()

        st.subheader("Editar Registro de Férias ou Abonada")
        if not df_folgas.empty:
            df_folgas['label'] = df_folgas.apply(lambda row: f"{row['tipo']} - {formatar_nome(row['nome_funcionario'])} ({pd.to_datetime(row['data_inicio']).strftime('%d/%m/%Y')})", axis=1)
            lista_eventos = ["Selecione um registro para editar..."] + df_folgas.sort_values(by='data_inicio', ascending=False)['label'].tolist()
            evento_label_selecionado = st.selectbox("Selecione o Registro", options=lista_eventos)

            if evento_label_selecionado != "Selecione um registro para editar...":
                evento_selecionado_df = df_folgas[df_folgas['label'] == evento_label_selecionado]
                if not evento_selecionado_df.empty:
                    dados_evento = evento_selecionado_df.iloc[0]
                    evento_id = dados_evento.name

                    with st.form(f"edit_folga_{evento_id}"):
                        st.write(f"Editando: **{dados_evento['label']}**")
                        tipo_evento_edit = dados_evento['tipo']
                        
                        if tipo_evento_edit == "Férias":
                            st.write("Período de Férias:")
                            col1_edit, col2_edit = st.columns(2)
                            with col1_edit:
                                data_inicio_edit = st.date_input("Nova Data de Início", value=pd.to_datetime(dados_evento['data_inicio']))
                            with col2_edit:
                                data_fim_edit = st.date_input("Nova Data de Fim", value=pd.to_datetime(dados_evento['data_fim']))
                        else:
                            st.write("Data da Abonada:")
                            data_inicio_edit = st.date_input("Nova Data", value=pd.to_datetime(dados_evento['data_inicio']))
                            data_fim_edit = data_inicio_edit

                        submit_edit = st.form_submit_button("Salvar Alterações")

                        if submit_edit:
                            if tipo_evento_edit == "Férias" and data_inicio_edit > data_fim_edit:
                                st.error("A data de início não pode ser posterior à data de fim.")
                            else:
                                try:
                                    ref = db.reference(f'folgas_ferias/{evento_id}')
                                    ref.update({'data_inicio': data_inicio_edit.strftime("%Y-%m-%d"),'data_fim': data_fim_edit.strftime("%Y-%m-%d")})
                                    
                                    log_atividade(st.session_state.get('username'), "Editou ausência", f"Registro: {dados_evento['label']}, Novas datas: {data_inicio_edit} a {data_fim_edit}")
                                    
                                    st.success("Registro atualizado com sucesso!")
                                    st.cache_data.clear()
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"Erro ao atualizar o registro: {e}")
                else:
                    st.warning("Registro não encontrado. Por favor, atualize a página.")
        else:
            st.info("Nenhum registro de férias ou abonada para editar.")
        st.divider()

        st.subheader("Histórico de Férias e Abonadas")
        if not df_folgas.empty:
            df_folgas_display = df_folgas.copy()
            df_folgas_display['nome_funcionario'] = df_folgas_display['nome_funcionario'].apply(formatar_nome)
            
            st.markdown("##### Filtrar Histórico")
            col1, col2, col3 = st.columns(3)
            with col1:
                funcionarios_disponiveis = sorted(df_folgas_display['nome_funcionario'].unique().tolist())
                filtro_funcionarios = st.multiselect("Filtrar por Funcionário(s)", options=funcionarios_disponiveis)
            with col2:
                filtro_tipo = st.selectbox("Filtrar por Tipo", ["Todos", "Férias", "Abonada"])
            with col3:
                df_folgas_display['ano'] = pd.to_datetime(df_folgas_display['data_inicio']).dt.year
                anos_disponiveis = sorted(df_folgas_display['ano'].unique(), reverse=True)
                filtro_ano = st.selectbox("Filtrar por Ano", ["Todos"] + anos_disponiveis)
                if filtro_funcionarios:
                    df_folgas_display = df_folgas_display[df_folgas_display['nome_funcionario'].isin(filtro_funcionarios)]
                if filtro_tipo != "Todos":
                    df_folgas_display = df_folgas_display[df_folgas_display['tipo'] == filtro_tipo]
                if filtro_ano != "Todos":
                    df_folgas_display = df_folgas_display[df_folgas_display['ano'] == filtro_ano]
            
            cols_to_display = ['nome_funcionario', 'tipo', 'data_inicio', 'data_fim']
            st.dataframe(df_folgas_display[cols_to_display].rename(columns={'nome_funcionario': 'Funcionário', 'tipo': 'Tipo', 'data_inicio': 'Início', 'data_fim': 'Fim'}), use_container_width=True,hide_index=True)
        else:
            st.write("Nenhum registro de ausência encontrado.")

    with tab_rh2:
        st.header("Visão Geral da Equipe")
        
        col_ficha, col_tabela = st.columns([0.7, 2.3])
        with col_tabela:
            st.subheader("Equipe e Status de Férias")
            if not df_funcionarios.empty and 'id' in df_funcionarios.columns:
                
                ferias_info_completa = [calcular_status_ferias_saldo(func, df_folgas) for _, func in df_funcionarios.iterrows()]
                
                df_display = df_funcionarios.copy()
                df_display['nome_formatado'] = df_display['nome'].apply(formatar_nome)
                df_display['Período Aquisitivo de Referência'] = [info[0] for info in ferias_info_completa]
                df_display['Status Agendamento'] = [info[1] for info in ferias_info_completa]
                df_display['status_code'] = [info[2] for info in ferias_info_completa] 
                df_display['Abonadas no Ano'] = [get_abonadas_ano(func_id, df_folgas) for func_id in df_funcionarios['id']]

                def style_status_code(code):
                    color = ''
                    if code == "PENDING": color = '#fff2cc'
                    elif code == "SCHEDULED": color = '#d4e6f1'
                    elif code == "ON_VACATION": color = '#d5f5e3'
                    elif code == "RISK_EXPIRING": color = '#f5b7b1'
                    return f'background-color: {color}'

                df_para_exibir = df_display[['nome_formatado', 'funcao', 'Período Aquisitivo de Referência', 'Status Agendamento', 'Abonadas no Ano']]
                df_renomeado = df_para_exibir.rename(columns={'nome_formatado': 'Nome', 'funcao': 'Função'})
                
                styler = df_renomeado.style.apply(
                    lambda row: [style_status_code(df_display.loc[row.name, 'status_code'])] * len(row),
                    axis=1
                )
                
                st.dataframe(
                    styler,
                    use_container_width=True,
                    hide_index=True
                )
            else:
                st.info("Nenhum funcionário cadastrado.")

        with col_ficha:
            st.subheader("Consultar Ficha")
            if lista_nomes_curtos:
                nome_curto_ficha = st.selectbox("Selecione um funcionário", lista_nomes_curtos, index=None, placeholder="Selecione...")
                if nome_curto_ficha:
                    nome_completo_ficha = nome_map[nome_curto_ficha]
                    dados_func = df_funcionarios[df_funcionarios['nome'] == nome_completo_ficha].iloc[0]
                    st.image("https://placehold.co/150x150/FFFFFF/333333?text=FOTO", use_container_width='auto')
                    st.markdown(f"**Nome:** {dados_func.get('nome', 'N/A')}")
                    st.markdown(f"**Matrícula:** {dados_func.get('matricula', 'N/A')}")
                    st.markdown(f"**Telefone:** {dados_func.get('telefone', 'N/A')}")
                    
                    data_adm_str = dados_func.get('data_admissao', 'N/A')
                    if data_adm_str != 'N/A':
                        data_adm_str = pd.to_datetime(data_adm_str).strftime('%d/%m/%Y')
                    st.markdown(f"**Data de Admissão:** {data_adm_str}")
                    
                    data_nasc_str = dados_func.get('data_nascimento')
                    if data_nasc_str:
                        try:
                            data_nasc_str = pd.to_datetime(data_nasc_str).strftime('%d/%m/%Y')
                        except (ValueError, TypeError):
                            data_nasc_str = "Data inválida"
                    else:
                        data_nasc_str = "N/A"

                    st.markdown(f"**Data de Nascimento:** {data_nasc_str}")
                    
                    st.markdown(f"**Tam. Camisa:** {dados_func.get('tamanho_camisa', 'N/A')}")
                    st.markdown(f"**Nº Bota:** {dados_func.get('numero_bota', 'N/A')}")
                    st.markdown(f"**Nº Chave:** {dados_func.get('numero_chave', 'N/A')}")

                    st.divider()
                    st.markdown("**Histórico Recente:**")

                    datas_abonadas = get_datas_abonadas_ano(dados_func.get('id'), df_folgas)
                    st.markdown(f"- **Abonadas no ano ({len(datas_abonadas)}):** {', '.join(datas_abonadas) if datas_abonadas else 'Nenhuma'}")
                    
                    ultimas_ferias = get_ultimas_ferias(dados_func.get('id'), df_folgas)
                    st.markdown(f"- **Últimas Férias:** {ultimas_ferias}")
            else:
                st.info("Nenhum funcionário.")

    with tab_rh3:
        st.subheader("Cadastrar Novo Funcionário")
        with st.form("novo_funcionario_form_3", clear_on_submit=True):
            nome = st.text_input("Nome Completo")
            matricula = st.text_input("Número da Matrícula")
            telefone = st.text_input("Telefone")
            funcao = st.text_input("Função")
            unidade_trabalho = st.text_input("Unidade de Trabalho")
            data_admissao = st.date_input("Data de Admissão", datetime.now())
            
            st.divider()
            st.markdown("**Informações Adicionais**")
            data_nascimento = st.date_input("Data de Nascimento", min_value=date(1940, 1, 1), max_value=date.today() - relativedelta(years=18), value=date.today() - relativedelta(years=25))
            
            col_uniforme1, col_uniforme2 = st.columns(2)
            with col_uniforme1:
                tamanho_camisa = st.text_input("Tamanho da Camisa (Ex: P, M, G, GG)")
            with col_uniforme2:
                numero_bota = st.text_input("Número da Bota (Ex: 40)")
            
            numero_chave = st.text_input("Número de Chave do Armário")
            
            submit_funcionario = st.form_submit_button("Cadastrar Funcionário")
            
            if submit_funcionario and nome and funcao and unidade_trabalho:
                try:
                    novo_id = str(int(time.time() * 1000))
                    ref = db.reference(f'funcionarios/{novo_id}')
                    dados_novos = {
                        'id': novo_id, 'nome': nome, 'matricula': matricula, 
                        'telefone': telefone, 'funcao': funcao, 'unidade_trabalho': unidade_trabalho, 
                        'data_admissao': data_admissao.strftime("%Y-%m-%d"),
                        'data_nascimento': data_nascimento.strftime("%Y-%m-%d"),
                        'tamanho_camisa': tamanho_camisa,
                        'numero_bota': numero_bota,
                        'numero_chave': numero_chave
                    }
                    ref.set(dados_novos)
                    
                    log_atividade(st.session_state.get('username'), "Cadastrou novo funcionário", f"Nome: {nome}")
                    
                    st.success(f"Funcionário {nome} cadastrado com sucesso!")
                    st.cache_data.clear(); st.rerun()
                except Exception as e:
                    st.error(f"Erro ao cadastrar funcionário: {e}")
        st.divider()
        st.subheader("Editar Funcionário")
        if lista_nomes_curtos:
            nome_curto_para_editar = st.selectbox("Selecione para editar", lista_nomes_curtos, index=None, placeholder="Selecione um funcionário...")
            if nome_curto_para_editar:
                nome_completo_para_editar = nome_map[nome_curto_para_editar]
                dados_func_originais = df_funcionarios[df_funcionarios['nome'] == nome_completo_para_editar].iloc[0]
                with st.form("edit_funcionario_form_3"):
                    st.write(f"Editando dados de **{nome_completo_para_editar}**")
                    nome_edit = st.text_input("Nome Completo", value=dados_func_originais.get('nome'))
                    matricula_edit = st.text_input("Número da Matrícula", value=dados_func_originais.get('matricula'))
                    telefone_edit = st.text_input("Telefone", value=dados_func_originais.get('telefone'))
                    funcao_edit = st.text_input("Função", value=dados_func_originais.get('funcao'))
                    unidade_edit = st.text_input("Unidade de Trabalho", value=dados_func_originais.get('unidade_trabalho'))
                    data_admissao_edit = st.date_input("Data de Admissão", value=pd.to_datetime(dados_func_originais.get('data_admissao')))

                    st.divider()
                    st.markdown("**Informações Adicionais**")
                    
                    data_nasc_val = pd.to_datetime(dados_func_originais.get('data_nascimento')) if pd.notna(dados_func_originais.get('data_nascimento')) else date.today() - relativedelta(years=25)
                    data_nascimento_edit = st.date_input("Data de Nascimento", value=data_nasc_val)
                    
                    col_edit_uniforme1, col_edit_uniforme2 = st.columns(2)
                    with col_edit_uniforme1:
                        tamanho_camisa_edit = st.text_input("Tamanho da Camisa", value=dados_func_originais.get('tamanho_camisa', ''))
                    with col_edit_uniforme2:
                        numero_bota_edit = st.text_input("Número da Bota", value=dados_func_originais.get('numero_bota', ''))
                    
                    numero_chave_edit = st.text_input("Número de Chave do Armário", value=dados_func_originais.get('numero_chave', ''))
                    
                    if st.form_submit_button("Salvar Alterações"):
                        dados_atualizados = {
                            'nome': nome_edit, 'matricula': matricula_edit, 
                            'telefone': telefone_edit, 'funcao': funcao_edit, 
                            'unidade_trabalho': unidade_edit, 
                            'data_admissao': data_admissao_edit.strftime('%Y-%m-%d'),
                            'data_nascimento': data_nascimento_edit.strftime('%Y-%m-%d'),
                            'tamanho_camisa': tamanho_camisa_edit,
                            'numero_bota': numero_bota_edit,
                            'numero_chave': numero_chave_edit
                        }
                        ref = db.reference(f"funcionarios/{dados_func_originais['id']}")
                        ref.update(dados_atualizados)
                        
                        log_atividade(st.session_state.get('username'), "Editou funcionário", f"Nome: {nome_edit}")
                        
                        st.success("Dados do funcionário atualizados com sucesso!")
                        st.cache_data.clear(); st.rerun()
        st.divider()
        st.subheader("🚨 Deletar Funcionário")
        if lista_nomes_curtos:
            nome_curto_para_deletar = st.selectbox("Selecione para deletar", lista_nomes_curtos, index=None, placeholder="Selecione um funcionário...")
            if nome_curto_para_deletar:
                nome_completo_para_deletar = nome_map[nome_curto_para_deletar]
                st.warning(f"**Atenção:** Você está prestes a deletar **{nome_completo_para_deletar}** e todos os seus registos. Esta ação é irreversível.")
                if st.button("Confirmar Deleção", type="primary"):
                    try:
                        id_func_deletar = df_funcionarios[df_funcionarios['nome'] == nome_completo_para_deletar]['id'].iloc[0]
                        db.reference(f'funcionarios/{id_func_deletar}').delete()
                        folgas_ref = db.reference('folgas_ferias')
                        folgas_para_deletar = folgas_ref.order_by_child('id_funcionario').equal_to(id_func_deletar).get()
                        if folgas_para_deletar:
                            for key in folgas_para_deletar:
                                folgas_ref.child(key).delete()
                        
                        log_atividade(st.session_state.get('username'), "Deletou funcionário", f"Nome: {nome_completo_para_deletar}")

                        st.success(f"Funcionário {nome_completo_para_deletar} deletado com sucesso.")
                        st.cache_data.clear(); st.rerun()
                    except Exception as e:
                        st.error(f"Ocorreu um erro ao deletar: {e}")

def modulo_denuncias():
    """Renderiza a pagina do modulo de Denuncias."""
    st.markdown("""
        <div class="mod-header">
            <h2>🚨 Denuncias</h2>
            <p>Registro, gerenciamento e acompanhamento de denuncias zoossanitarias</p>
        </div>
    """, unsafe_allow_html=True)

    # Lista fixa de motivos para padronização
    lista_motivos_denuncia = [
        "Acúmulo de lixo/entulho",
        "Maus tratos a animais",
        "Animal de grande porte em via pública",
        "Foco de dengue/escorpião",
        "Criação inadequada de animais",
        "Comércio irregular de animais",
        "Outros"
    ]

    # Carrega funcionários para usar na lista de responsáveis
    df_funcionarios = carregar_dados_firebase('funcionarios')
    if not df_funcionarios.empty:
        lista_responsaveis = sorted([formatar_nome(nome) for nome in df_funcionarios['nome']])
    else:
        lista_responsaveis = []
        
    @st.cache_data
    def geocode_addresses(df):
        geolocator = Nominatim(user_agent=f"streamlit_app_{int(time.time())}")
        latitudes, longitudes = [], []
        df_copy = df.copy()
        for col in ['logradouro', 'numero', 'bairro', 'cep']:
            if col not in df_copy.columns: df_copy[col] = ''
        for index, row in df_copy.iterrows():
            address = f"{row.get('logradouro', '')}, {row.get('numero', '')}, {row.get('bairro', '')}, Guaratinguetá, SP, Brasil"
            try:
                location = geolocator.geocode(address, timeout=10)
                if location:
                    latitudes.append(location.latitude)
                    longitudes.append(location.longitude)
                else:
                    latitudes.append(None)
                    longitudes.append(None)
                time.sleep(1) # Para evitar sobrecarregar o serviço de geocodificação
            except Exception as e:
                latitudes.append(None)
                longitudes.append(None)
        df_copy['lat'], df_copy['lon'] = latitudes, longitudes
        return df_copy.dropna(subset=['lat', 'lon'])

    def carregar_e_cachear_denuncias():
        ref = db.reference('denuncias')
        denuncias_data = ref.get()
        if denuncias_data:
            denuncias_padronizadas = []
            for protocolo, dados in denuncias_data.items():
                if isinstance(dados, dict):
                    dados['protocolo'] = protocolo
                    dados.setdefault('logradouro', dados.get('rua', ''))
                    dados.setdefault('conclusao_atendimento', '')
                    dados.setdefault('cep', '')
                    dados.setdefault('status', 'Não atendida')
                    dados.setdefault('auto_infracao', 'Não')
                    dados.setdefault('protocolo_auto_infracao', '')
                    dados.setdefault('auto_imposicao_penalidade', 'Não')
                    dados.setdefault('protocolo_auto_imposicao_penalidade', '')
                    dados.setdefault('responsavel_atendimento', '')
                    dados.setdefault('relatorio_atendimento', '')
                    # Adiciona os novos campos com valores padrão para evitar erros
                    dados.setdefault('data_atendimento', None)
                    dados.setdefault('responsavel_imovel', '')
                    dados.setdefault('rg_responsavel', '')
                    dados.setdefault('cpf_responsavel', '')
                    denuncias_padronizadas.append(dados)
            df = pd.DataFrame(denuncias_padronizadas)
            if 'protocolo' in df.columns:
                df['protocolo_int'] = df['protocolo'].apply(lambda x: int(x) if str(x).isdigit() else 0)
                df = df.sort_values(by='protocolo_int', ascending=False)
                del df['protocolo_int']
            st.session_state.denuncias_df = df
        else:
            st.session_state.denuncias_df = pd.DataFrame()

    if 'denuncias_df' not in st.session_state:
        carregar_e_cachear_denuncias()
    
    tab1, tab2, tab3 = st.tabs(["📋 Registrar Denúncia", "🛠️ Gerenciamento", "📊 Dashboard"])
    
    with tab1:
        st.subheader("Registrar Nova Denúncia")
        with st.form("nova_denuncia_form", clear_on_submit=True):
            data_denuncia = st.date_input("Data da Denúncia", datetime.now())
            # Campo de motivo alterado para selectbox
            motivo_denuncia = st.selectbox("Motivo da Denúncia", options=lista_motivos_denuncia)
            
            bairro = st.text_input("Bairro")
            logradouro = st.text_input("Logradouro")
            numero = st.text_input("Nº")
            cep = st.text_input("CEP (Opcional)")
            detalhes_denuncia = st.text_area("Detalhes da Denúncia")
            submit_button = st.form_submit_button("Registrar Denúncia")
            
        if submit_button:
            if motivo_denuncia and logradouro and bairro:
                ano_atual = datetime.now().year
                ref_contador = db.reference(f'contadores/{ano_atual}')
                def incrementar(valor_atual):
                    return (valor_atual or 0) + 1
                
                novo_numero = ref_contador.transaction(incrementar)
                protocolo_gerado = f"{novo_numero:04d}{ano_atual}"

                if protocolo_gerado:
                    nova_denuncia = {
                        "data_denuncia": data_denuncia.strftime("%Y-%m-%d"), "motivo_denuncia": motivo_denuncia, 
                        "bairro": bairro, "logradouro": logradouro, "numero": numero, "cep": cep, 
                        "detalhes_denuncia": detalhes_denuncia, "status": "Não atendida", 
                        "auto_infracao": "Não", "protocolo_auto_infracao": "", 
                        "auto_imposicao_penalidade": "Não", "protocolo_auto_imposicao_penalidade": "", 
                        "responsavel_atendimento": "", "relatorio_atendimento": "", "conclusao_atendimento": "",
                        "data_atendimento": None, "responsavel_imovel": "", "rg_responsavel": "", "cpf_responsavel": ""
                    }
                    ref = db.reference(f'denuncias/{protocolo_gerado}')
                    ref.set(nova_denuncia)
                    
                    log_atividade(st.session_state.get('username'), "Registrou nova denúncia", f"Protocolo: {protocolo_gerado}")

                    st.success(f"Denúncia registrada com sucesso! Protocolo: {protocolo_gerado}")
                    carregar_e_cachear_denuncias()
                    st.cache_data.clear()
                    st.rerun()
            else:
                st.warning("Por favor, preencha os campos obrigatórios (Motivo, Bairro, Logradouro).")
        st.divider()
        st.subheader("Denúncias Recentes")
        if 'denuncias_df' in st.session_state and not st.session_state.denuncias_df.empty:
            cols = ['protocolo', 'data_denuncia', 'motivo_denuncia', 'bairro', 'logradouro', 'numero']
            df_display = st.session_state.denuncias_df[[c for c in cols if c in st.session_state.denuncias_df.columns]]
            df_display = df_display.rename(columns={'protocolo': 'PROTOCOLO','data_denuncia': 'DATA','motivo_denuncia': 'MOTIVO','bairro': 'BAIRRO','logradouro': 'LOGRADOURO','numero': 'Nº'})
            st.dataframe(df_display,hide_index=True,use_container_width=True)

    with tab2:
        if 'denuncias_df' in st.session_state and not st.session_state.denuncias_df.empty:
            protocolo_selecionado = st.selectbox("Selecione o Protocolo para Gerenciar", options=st.session_state.denuncias_df['protocolo'].tolist(), index=None, placeholder="Selecione um protocolo...")
            if protocolo_selecionado:
                dados_denuncia = st.session_state.denuncias_df[st.session_state.denuncias_df['protocolo'] == protocolo_selecionado].iloc[0]
                with st.form("gerenciamento_form"):
                    st.subheader(f"Atualizando Protocolo: {protocolo_selecionado}")
                    
                    status = st.selectbox("Status", options=["Não atendida", "Atendida", "Arquivada"], index=["Não atendida", "Atendida", "Arquivada"].index(dados_denuncia.get('status', 'Não atendida')))
                    
                    # Responsável pelo atendimento como lista
                    responsavel_atendimento = st.selectbox("Responsável pelo Atendimento", options=[""] + lista_responsaveis, index=lista_responsaveis.index(dados_denuncia.get('responsavel_atendimento')) + 1 if dados_denuncia.get('responsavel_atendimento') in lista_responsaveis else 0)

                    # Data do atendimento
                    data_atendimento_val = pd.to_datetime(dados_denuncia.get('data_atendimento')).date() if dados_denuncia.get('data_atendimento') else None
                    data_atendimento = st.date_input("Data do Atendimento", value=data_atendimento_val)

                    # Cálculo e exibição da data de retorno
                    if data_atendimento:
                        data_retorno = data_atendimento + timedelta(days=14)
                        st.info(f"ℹ️ Data de Retorno: {data_retorno.strftime('%d/%m/%Y')}")

                    st.divider()
                    st.markdown("**Dados do Responsável pelo Imóvel**")
                    responsavel_imovel = st.text_input("Nome do Responsável do Imóvel", value=dados_denuncia.get('responsavel_imovel', ''))
                    col_doc1, col_doc2 = st.columns(2)
                    with col_doc1:
                        rg_responsavel = st.text_input("RG (Opcional)", value=dados_denuncia.get('rg_responsavel', ''))
                    with col_doc2:
                        cpf_responsavel = st.text_input("CPF (Opcional)", value=dados_denuncia.get('cpf_responsavel', ''))
                    
                    st.divider()
                    st.markdown("**Relatório e Conclusão**")
                    relatorio = st.text_area("Relatório (Situação Encontrada)", value=dados_denuncia.get('relatorio_atendimento', ''), height=150)
                    conclusao = st.text_area("Conclusão do Atendimento", value=dados_denuncia.get('conclusao_atendimento', ''), height=150)
                    
                    st.divider()
                    col1, col2 = st.columns(2)
                    with col1:
                        auto_infracao = st.selectbox("Auto de Infração?", options=["Não", "Sim"], index=["Não", "Sim"].index(dados_denuncia.get('auto_infracao', 'Não')))
                        protocolo_auto_infracao = st.text_input("Nº Auto de Infração", value=dados_denuncia.get('protocolo_auto_infracao', '')) if auto_infracao == "Sim" else ""
                    with col2:
                        auto_penalidade = st.selectbox("Auto de Penalidade?", options=["Não", "Sim"], index=["Não", "Sim"].index(dados_denuncia.get('auto_imposicao_penalidade', 'Não')))
                        protocolo_auto_penalidade = st.text_input("Nº Auto de Penalidade", value=dados_denuncia.get('protocolo_auto_imposicao_penalidade', '')) if auto_penalidade == "Sim" else ""

                    if st.form_submit_button("Salvar Gerenciamento"):
                        if not responsavel_imovel:
                               st.error("O campo 'Nome do Responsável do Imóvel' é de preenchimento obrigatório.")
                        else:
                            dados_para_atualizar = {
                                "status": status, 
                                "responsavel_atendimento": responsavel_atendimento, 
                                "data_atendimento": data_atendimento.strftime("%Y-%m-%d") if data_atendimento else None,
                                "responsavel_imovel": responsavel_imovel,
                                "rg_responsavel": rg_responsavel,
                                "cpf_responsavel": cpf_responsavel,
                                "relatorio_atendimento": relatorio, 
                                "conclusao_atendimento": conclusao, 
                                "auto_infracao": auto_infracao, 
                                "protocolo_auto_infracao": protocolo_auto_infracao, 
                                "auto_imposicao_penalidade": auto_penalidade, 
                                "protocolo_auto_imposicao_penalidade": protocolo_auto_penalidade
                            }
                            ref = db.reference(f'denuncias/{protocolo_selecionado}')
                            ref.update(dados_para_atualizar)
                            
                            log_atividade(st.session_state.get('username'), "Atualizou denúncia", f"Protocolo: {protocolo_selecionado}, Status: {status}")

                            st.success(f"Denúncia {protocolo_selecionado} atualizada!")
                            carregar_e_cachear_denuncias()
                            st.cache_data.clear()
                            st.rerun()

                with st.expander("🚨 Deletar Denúncia"):
                    if st.button("Eu entendo o risco, deletar denúncia", type="primary"):
                        ref = db.reference(f'denuncias/{protocolo_selecionado}'); ref.delete()
                        
                        log_atividade(st.session_state.get('username'), "Deletou denúncia", f"Protocolo: {protocolo_selecionado}")

                        st.success(f"Denúncia {protocolo_selecionado} deletada!")
                        carregar_e_cachear_denuncias(); st.cache_data.clear(); st.rerun()
        else:
            st.info("Nenhuma denúncia registrada para gerenciar.")

    with tab3:
        if 'denuncias_df' in st.session_state and not st.session_state.denuncias_df.empty:
            df_resumo = st.session_state.denuncias_df.copy()
            st.subheader("Métricas Gerais"); status_counts = df_resumo['status'].value_counts()
            col1, col2, col3, col4 = st.columns(4)
            col1.metric("Denúncias Totais", len(df_resumo)); col2.metric("Atendidas", status_counts.get('Atendida', 0))
            col3.metric("Não Atendidas", status_counts.get('Não atendida', 0)); col4.metric("Arquivadas", status_counts.get('Arquivada', 0))
            st.divider()
            st.subheader("Gerar Relatório de Denúncia (.docx)")
            protocolo_relatorio = st.selectbox("Selecione um Protocolo para gerar relatório", options=df_resumo['protocolo'].tolist(), index=None, placeholder="Escolha o protocolo...")
            if protocolo_relatorio:
                dados_relatorio = df_resumo[df_resumo['protocolo'] == protocolo_relatorio].iloc[0]
                report_bytes = create_word_report(dados_relatorio)
                st.download_button(label="📥 Baixar Relatório em Word", data=report_bytes, file_name=f"Relatorio_Inspecao_{protocolo_relatorio}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            st.divider()
            st.subheader("Tabela de Resumo")
            cols_resumo = ['protocolo', 'data_denuncia', 'motivo_denuncia', 'status', 'responsavel_atendimento', 'data_atendimento', 'responsavel_imovel']
            df_resumo_display = df_resumo[[c for c in cols_resumo if c in df_resumo.columns]]
            st.dataframe(df_resumo_display.rename(columns={'protocolo': 'Protocolo', 'data_denuncia': 'Data Denúncia', 'motivo_denuncia': 'Motivo', 'status': 'Status', 'responsavel_atendimento': 'Resp. Atendimento', 'data_atendimento': 'Data Atendimento', 'responsavel_imovel': 'Resp. Imóvel'}), use_container_width=True)
            st.divider()
            st.subheader("Análise Gráfica")
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("##### Denúncias Atendidas por Mês")
                df_atendidas = df_resumo[df_resumo['status'] == 'Atendida'].copy()
                if not df_atendidas.empty and 'data_denuncia' in df_atendidas:
                    df_atendidas['data_denuncia'] = pd.to_datetime(df_atendidas['data_denuncia'])
                    df_atendidas['mes_ano'] = df_atendidas['data_denuncia'].dt.to_period('M').astype(str)
                    atendidas_por_mes = df_atendidas['mes_ano'].value_counts().sort_index()
                    fig_bar = px.bar(atendidas_por_mes, x=atendidas_por_mes.index, y=atendidas_por_mes.values, title="Contagem de Denúncias Atendidas Mensalmente", labels={'x': 'Mês/Ano', 'y': 'Quantidade'}, text_auto=True)
                    st.plotly_chart(fig_bar, use_container_width=True)
                else:
                    st.info("Nenhuma denúncia 'Atendida' para exibir no gráfico.")
            with col2:
                st.markdown("##### Distribuição de Denúncias por Motivo")
                denuncias_por_motivo = df_resumo['motivo_denuncia'].value_counts()
                fig_pie = px.pie(denuncias_por_motivo, values=denuncias_por_motivo.values, names=denuncias_por_motivo.index, title="Distribuição por Motivo", hole=.3, color_discrete_sequence=px.colors.sequential.RdBu)
                st.plotly_chart(fig_pie, use_container_width=True)
            st.divider()
            st.subheader("Geolocalização das Denúncias")
            with st.spinner("Geocodificando endereços..."):
                df_mapeado = geocode_addresses(df_resumo)
            if not df_mapeado.empty:
                st.map(df_mapeado, latitude='lat', longitude='lon', size=10)
            else:
                st.warning("Não foi possível geolocalizar nenhum endereço.")
        else:
            st.info("Nenhuma denúncia registrada.")

# --- MÓDULO DO BOLETIM (LAYOUT CORRIGIDO) ---
def modulo_boletim():
    """Renderiza a pagina do modulo de Boletim de Programacao Diaria com um layout melhorado."""
    st.markdown("""
        <div class="mod-header">
            <h2>🗓️ Boletim de Programacao Diaria</h2>
            <p>Programacao de equipes, atividades e mapa de campo</p>
        </div>
    """, unsafe_allow_html=True)

    # CSS antigo removido - agora usa CSS global


    # Carregamento dos dados necessários
    df_funcionarios = carregar_dados_firebase('funcionarios')
    df_boletins = carregar_dados_firebase('boletins')
    lista_quarteiroes = carregar_quarteiroes_csv()
    df_geo_quarteiroes = carregar_geo_kml()

    # Controle do estado para equipes dinâmicas
    if 'num_equipes_manha' not in st.session_state:
        st.session_state.num_equipes_manha = 1
    if 'num_equipes_tarde' not in st.session_state:
        st.session_state.num_equipes_tarde = 1
    if 'num_equipes_pe_ie' not in st.session_state:
        st.session_state.num_equipes_pe_ie = 1

    tab1, tab_pe_ie, tab2, tab3, tab4 = st.tabs(["🗓️ Criar Boletim", "📍 P.E e I.E", "🔍 Visualizar/Editar Boletim", "🗺️ Mapa de Atividades", "📊 Dashboard"])

    with tab1:
        st.header("Novo Boletim de Programação")

        col1, col2 = st.columns(2)

        with col1:
            # Card: Dados do Boletim
            st.markdown("<div class='sys-card'>", unsafe_allow_html=True)
            st.markdown("<div class='sys-card-title'>Dados do Boletim</div>", unsafe_allow_html=True)

            data_boletim = st.date_input("Data do Trabalho", date.today())
            bairros = st.text_area("Bairros a serem trabalhados")
            atividades_gerais_options = ["Controle de criadouros", "Visita a Imóveis", "ADL", "Nebulização"]
            atividades_gerais = st.multiselect("Atividades Gerais do Dia", atividades_gerais_options)

            if isinstance(df_funcionarios, pd.DataFrame) and not df_funcionarios.empty:
                nome_map = {formatar_nome(nome): nome for nome in df_funcionarios['nome']}
                lista_nomes_curtos_full = sorted(list(nome_map.keys()))
            else:
                nome_map = {}
                lista_nomes_curtos_full = []
                st.warning("Não há funcionários cadastrados para criar um boletim.")

            motoristas_curtos = st.multiselect("Motorista(s)", options=lista_nomes_curtos_full)
            st.markdown("</div>", unsafe_allow_html=True) # Fim do card

            # Card: Equipes
            st.markdown("<div class='sys-card'>", unsafe_allow_html=True)
            st.markdown("<div class='sys-card-title'>Equipes</div>", unsafe_allow_html=True)
            
            col_equipe_manhã, col_equipe_tarde = st.columns(2)
            
            with col_equipe_manhã:
                st.markdown("#### Manhã")
                equipes_manha = []
                membros_selecionados_manha = []
                nomes_disponiveis_manha = [nome for nome in lista_nomes_curtos_full]
                
                # Removendo ausentes e motoristas da lista
                if 'faltas_manha_curtos' in st.session_state and st.session_state.faltas_manha_curtos is not None:
                    nomes_disponiveis_manha = [nome for nome in nomes_disponiveis_manha if nome not in st.session_state.faltas_manha_curtos]
                if motoristas_curtos:
                    nomes_disponiveis_manha = [nome for nome in nomes_disponiveis_manha if nome not in motoristas_curtos]

                for i in range(st.session_state.num_equipes_manha):
                    if i > 0:
                        st.markdown("<hr>", unsafe_allow_html=True)
                        
                    st.markdown(f"**Equipe {i+1}**")
                    opcoes_equipe_manha = [nome for nome in nomes_disponiveis_manha if nome not in membros_selecionados_manha]
                    
                    membros_curtos = st.multiselect("Membros", options=opcoes_equipe_manha, key=f"manha_membros_{i}")
                    atividades = st.multiselect("Atividades", options=atividades_gerais_options, key=f"manha_atividades_{i}")
                    quarteiroes = st.multiselect("Quarteirões", options=lista_quarteiroes, key=f"manha_quarteiroes_{i}")
                    
                    if membros_curtos:
                        membros_completos = [nome_map[nome] for nome in membros_curtos]
                        equipes_manha.append({"membros": membros_completos, "atividades": atividades, "quarteiroes": quarteiroes})
                        membros_selecionados_manha.extend(membros_curtos)

                if st.button("➕ Adicionar Equipe (Manhã)", key="add_equipe_manha_button"):
                    st.session_state.num_equipes_manha += 1
                    st.rerun()

            with col_equipe_tarde:
                st.markdown("#### Tarde")
                equipes_tarde = []
                membros_selecionados_tarde = []
                nomes_disponiveis_tarde = [nome for nome in lista_nomes_curtos_full]
                
                # Removendo ausentes e motoristas da lista
                if 'faltas_tarde_curtos' in st.session_state and st.session_state.faltas_tarde_curtos is not None:
                    nomes_disponiveis_tarde = [nome for nome in nomes_disponiveis_tarde if nome not in st.session_state.faltas_tarde_curtos]
                if motoristas_curtos:
                    nomes_disponiveis_tarde = [nome for nome in nomes_disponiveis_tarde if nome not in motoristas_curtos]

                for i in range(st.session_state.num_equipes_tarde):
                    if i > 0:
                        st.markdown("<hr>", unsafe_allow_html=True)

                    st.markdown(f"**Equipe {i+1}**")
                    opcoes_equipe_tarde = [nome for nome in nomes_disponiveis_tarde if nome not in membros_selecionados_tarde]

                    membros_curtos = st.multiselect("Membros ", options=opcoes_equipe_tarde, key=f"tarde_membros_{i}")
                    atividades = st.multiselect("Atividades ", options=atividades_gerais_options, key=f"tarde_atividades_{i}")
                    quarteiroes = st.multiselect("Quarteirões ", options=lista_quarteiroes, key=f"tarde_quarteiroes_{i}")
                    
                    if membros_curtos:
                        membros_completos = [nome_map[nome] for nome in membros_curtos]
                        equipes_tarde.append({"membros": membros_completos, "atividades": atividades, "quarteiroes": quarteiroes})
                        membros_selecionados_tarde.extend(membros_curtos)

                if st.button("➕ Adicionar Equipe (Tarde)", key="add_equipe_tarde_button"):
                    st.session_state.num_equipes_tarde += 1
                    st.rerun()

            st.markdown("</div>", unsafe_allow_html=True) # Fim do card

        with col2:
            # Card: Ausências do Dia
            st.markdown("<div class='sys-card'>", unsafe_allow_html=True)
            st.markdown("<div class='sys-card-title'>Ausências do Dia</div>", unsafe_allow_html=True)
            
            ausencias_col1, ausencias_col2 = st.columns(2)
            with ausencias_col1:
                st.subheader("Manhã")
                faltas_manha_curtos = st.multiselect("Ausentes", options=lista_nomes_curtos_full, key="faltas_manha_curtos")
                motivo_falta_manha = st.text_input("Motivo", key="motivo_falta_manha")
            with ausencias_col2:
                st.subheader("Tarde")
                faltas_tarde_curtos = st.multiselect("Ausentes", options=lista_nomes_curtos_full, key="faltas_tarde_curtos")
                motivo_falta_tarde = st.text_input("Motivo", key="motivo_falta_tarde")
            
            st.markdown("</div>", unsafe_allow_html=True) # Fim do card
            
        # Botão de salvar no final
        if st.button("Salvar Boletim", use_container_width=True, type="primary", key="save_boletim_button"):
            motoristas_completos = [nome_map[nome] for nome in motoristas_curtos]
            faltas_manha_completos = [nome_map[nome] for nome in faltas_manha_curtos]
            faltas_tarde_completos = [nome_map[nome] for nome in faltas_tarde_curtos]
            
            boletim_id = data_boletim.strftime("%Y-%m-%d")
            boletim_data = {
                "data": boletim_id,
                "bairros": bairros,
                "atividades_gerais": atividades_gerais,
                "motoristas": motoristas_completos,
                "equipes_manha": equipes_manha,
                "equipes_tarde": equipes_tarde,
                "faltas_manha": {"nomes": faltas_manha_completos, "motivo": motivo_falta_manha},
                "faltas_tarde": {"nomes": faltas_tarde_completos, "motivo": motivo_falta_tarde}
            }
            try:
                ref = db.reference(f'boletins/{boletim_id}'); ref.set(boletim_data)
                
                log_atividade(st.session_state.get('username'), "Criou novo boletim", f"Data: {boletim_id}, Bairros: {bairros}")

                st.success(f"Boletim para o dia {data_boletim.strftime('%d/%m/%Y')} salvo com sucesso!")
                st.cache_data.clear()
                time.sleep(1)
                st.rerun()
            except Exception as e:
                st.error(f"Erro ao salvar o boletim: {e}")


    with tab_pe_ie:

        # Header
        st.markdown("""
            <div class="mod-header">
                <h2>Pontos Estrategicos e Imoveis Especiais</h2>
                <p>P.E = Ponto de Encontro (quinzenal) &nbsp;&middot;&nbsp; I.E = Imovel Especial (trimestral)</p>
            </div>
        """, unsafe_allow_html=True)

        df_pe_ie = carregar_dados_firebase('pe_ie_cadastros')
        df_boletins_pe_ie = carregar_dados_firebase('boletins_pe_ie')

        if isinstance(df_funcionarios, pd.DataFrame) and not df_funcionarios.empty:
            nome_map_pe = {formatar_nome(nome): nome for nome in df_funcionarios['nome']}
            lista_nomes_pe = sorted(list(nome_map_pe.keys()))
        else:
            nome_map_pe = {}
            lista_nomes_pe = []

        lista_pe_ie_opcoes = []
        if not df_pe_ie.empty:
            for idx_pe, row_pe in df_pe_ie.iterrows():
                label = f"{row_pe.get('tipo', '')} - {row_pe.get('nome_fantasia', '')} (No {row_pe.get('numero_cadastro', '')})"
                lista_pe_ie_opcoes.append({"id": idx_pe, "label": label, "tipo": row_pe.get('tipo', ''), "dados": row_pe})

        # Metricas resumo
        total_pe = len(df_pe_ie[df_pe_ie['tipo'] == 'P.E']) if not df_pe_ie.empty and 'tipo' in df_pe_ie.columns else 0
        total_ie = len(df_pe_ie[df_pe_ie['tipo'] == 'I.E']) if not df_pe_ie.empty and 'tipo' in df_pe_ie.columns else 0
        total_boletins = len(df_boletins_pe_ie) if not df_boletins_pe_ie.empty else 0

        st.markdown(f"""
            <div class="metric-row">
                <div class="metric-box">
                    <div class="metric-number">{total_pe}</div>
                    <div class="metric-label">Pontos de Encontro</div>
                </div>
                <div class="metric-box">
                    <div class="metric-number">{total_ie}</div>
                    <div class="metric-label">Imoveis Especiais</div>
                </div>
                <div class="metric-box">
                    <div class="metric-number">{total_pe + total_ie}</div>
                    <div class="metric-label">Total Cadastrados</div>
                </div>
                <div class="metric-box">
                    <div class="metric-number">{total_boletins}</div>
                    <div class="metric-label">Boletins Registrados</div>
                </div>
            </div>
        """, unsafe_allow_html=True)

        sub_tab_cadastrar, sub_tab_boletim, sub_tab_listar, sub_tab_historico = st.tabs([
            "📝 Cadastrar Imovel",
            "🗓️ Criar Boletim",
            "📋 Imoveis Cadastrados",
            "🔍 Historico de Boletins"
        ])

        # =============================================
        # SUB-ABA 1: CADASTRAR P.E / I.E
        # =============================================
        with sub_tab_cadastrar:

            st.markdown('<div class="sys-card"><div class="sys-card-title">📋 Dados do Novo Imovel</div>', unsafe_allow_html=True)

            tipo_pe_ie = st.selectbox("Tipo do Imovel", ["P.E - Ponto de Encontro", "I.E - Imovel Especial"], key="tipo_pe_ie_cadastro")

            col_cad1, col_cad2 = st.columns(2)
            with col_cad1:
                numero_cadastro = st.text_input("Numero de Cadastro", key="num_cadastro_pe")
                nome_fantasia = st.text_input("Nome Fantasia do Imovel", key="nome_fantasia_pe")
                endereco_pe = st.text_input("Endereco Completo", key="endereco_pe")
            with col_cad2:
                quarteirao_pe = st.selectbox("Numero de Quarteirao", options=[""] + lista_quarteiroes, key="quarteirao_pe")
                col_coord1, col_coord2 = st.columns(2)
                with col_coord1:
                    latitude_pe = st.text_input("Latitude", key="lat_pe", placeholder="-22.8136")
                with col_coord2:
                    longitude_pe = st.text_input("Longitude", key="lon_pe", placeholder="-45.1917")

            st.markdown('</div>', unsafe_allow_html=True)

            if st.button("✅ Salvar Cadastro", use_container_width=True, type="primary", key="save_pe_ie_button"):
                if numero_cadastro and endereco_pe and nome_fantasia:
                    tipo_sigla = "P.E" if "P.E" in tipo_pe_ie else "I.E"
                    frequencia = "Quinzenal" if tipo_sigla == "P.E" else "Trimestral"

                    cadastro_id = str(int(time.time() * 1000))
                    cadastro_data = {
                        "tipo": tipo_sigla,
                        "frequencia": frequencia,
                        "numero_cadastro": numero_cadastro,
                        "endereco": endereco_pe,
                        "nome_fantasia": nome_fantasia,
                        "latitude": latitude_pe,
                        "longitude": longitude_pe,
                        "quarteirao": quarteirao_pe,
                        "data_cadastro": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    }
                    try:
                        ref = db.reference(f'pe_ie_cadastros/{cadastro_id}')
                        ref.set(cadastro_data)

                        log_atividade(st.session_state.get('username'), f"Cadastrou {tipo_sigla}", f"No Cadastro: {numero_cadastro}, Nome: {nome_fantasia}")

                        st.success(f"{tipo_sigla} — {nome_fantasia} cadastrado com sucesso!")
                        st.cache_data.clear()
                        st.rerun()
                    except Exception as e:
                        st.error(f"Erro ao salvar o cadastro: {e}")
                else:
                    st.warning("Preencha os campos obrigatorios: Numero de Cadastro, Endereco e Nome Fantasia.")

        # =============================================
        # SUB-ABA 2: CRIAR BOLETIM DE P.E / I.E
        # =============================================
        with sub_tab_boletim:

            if not lista_pe_ie_opcoes:
                st.markdown("""
                    <div class="empty-state">
                        <div class="icon">📭</div>
                        <p>Nenhum P.E ou I.E cadastrado ainda.<br>Cadastre imoveis na aba <strong>"Cadastrar Imovel"</strong> para criar boletins.</p>
                    </div>
                """, unsafe_allow_html=True)
            else:
                col_bol1, col_bol2 = st.columns(2)

                with col_bol1:
                    st.markdown('<div class="sys-card"><div class="sys-card-title">📄 Dados do Boletim</div>', unsafe_allow_html=True)

                    data_boletim_pe = st.date_input("Data do Trabalho", date.today(), key="data_boletim_pe_ie")

                    filtro_tipo_boletim = st.selectbox("Filtrar imoveis por tipo", ["Todos", "P.E", "I.E"], key="filtro_tipo_boletim_pe")

                    opcoes_imoveis = [item["label"] for item in lista_pe_ie_opcoes if filtro_tipo_boletim == "Todos" or item["tipo"] == filtro_tipo_boletim]

                    imoveis_selecionados = st.multiselect(
                        "Selecione os imoveis trabalhados",
                        options=opcoes_imoveis,
                        key="imoveis_selecionados_pe_ie"
                    )

                    observacoes_pe = st.text_area("Observacoes gerais", key="obs_boletim_pe_ie", height=100)

                    st.markdown('</div>', unsafe_allow_html=True)

                with col_bol2:
                    st.markdown('<div class="sys-card"><div class="sys-card-title">👥 Equipes</div>', unsafe_allow_html=True)

                    equipes_pe_ie = []
                    membros_selecionados_pe = []

                    for i in range(st.session_state.num_equipes_pe_ie):
                        if i > 0:
                            st.markdown('<div class="sys-divider"></div>', unsafe_allow_html=True)

                        st.markdown(f"**Equipe {i+1}**")
                        opcoes_equipe_pe = [nome for nome in lista_nomes_pe if nome not in membros_selecionados_pe]

                        membros_pe_curtos = st.multiselect("Membros", options=opcoes_equipe_pe, key=f"pe_ie_membros_{i}")

                        if membros_pe_curtos:
                            membros_pe_completos = [nome_map_pe[nome] for nome in membros_pe_curtos]
                            equipes_pe_ie.append({"membros": membros_pe_completos})
                            membros_selecionados_pe.extend(membros_pe_curtos)

                    if st.button("➕ Adicionar Equipe", key="add_equipe_pe_ie_button"):
                        st.session_state.num_equipes_pe_ie += 1
                        st.rerun()

                    st.markdown('</div>', unsafe_allow_html=True)

                st.markdown("")
                if st.button("✅ Salvar Boletim P.E / I.E", use_container_width=True, type="primary", key="save_boletim_pe_ie"):
                    if imoveis_selecionados and equipes_pe_ie:
                        boletim_pe_id = str(int(time.time() * 1000))
                        boletim_pe_data = {
                            "data": data_boletim_pe.strftime("%Y-%m-%d"),
                            "imoveis_trabalhados": imoveis_selecionados,
                            "equipes": equipes_pe_ie,
                            "observacoes": observacoes_pe,
                            "criado_por": st.session_state.get('username', ''),
                            "data_criacao": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        }
                        try:
                            ref = db.reference(f'boletins_pe_ie/{boletim_pe_id}')
                            ref.set(boletim_pe_data)

                            log_atividade(
                                st.session_state.get('username'),
                                "Criou boletim P.E/I.E",
                                f"Data: {data_boletim_pe.strftime('%d/%m/%Y')}, Imoveis: {len(imoveis_selecionados)}"
                            )

                            st.success(f"Boletim P.E/I.E para {data_boletim_pe.strftime('%d/%m/%Y')} salvo com sucesso!")
                            st.session_state.num_equipes_pe_ie = 1
                            st.cache_data.clear()
                            time.sleep(1)
                            st.rerun()
                        except Exception as e:
                            st.error(f"Erro ao salvar o boletim: {e}")
                    else:
                        st.warning("Selecione pelo menos um imovel e monte pelo menos uma equipe.")

        # =============================================
        # SUB-ABA 3: CADASTROS EXISTENTES
        # =============================================
        with sub_tab_listar:

            if not df_pe_ie.empty:
                col_filtro1, col_filtro2 = st.columns(2)
                with col_filtro1:
                    filtro_tipo_pe = st.selectbox("Filtrar por tipo", ["Todos", "P.E", "I.E"], key="filtro_tipo_pe_ie")
                with col_filtro2:
                    busca_nome = st.text_input("🔍 Buscar por nome fantasia", key="busca_nome_pe_ie")

                df_pe_ie_display = df_pe_ie.copy()

                if filtro_tipo_pe != "Todos":
                    df_pe_ie_display = df_pe_ie_display[df_pe_ie_display['tipo'] == filtro_tipo_pe]

                if busca_nome:
                    df_pe_ie_display = df_pe_ie_display[df_pe_ie_display['nome_fantasia'].str.contains(busca_nome, case=False, na=False)]

                if df_pe_ie_display.empty:
                    st.info("Nenhum cadastro encontrado com os filtros aplicados.")
                else:
                    for idx, cadastro in df_pe_ie_display.iterrows():
                        tipo_label = cadastro.get('tipo', 'N/A')
                        freq_label = cadastro.get('frequencia', '')
                        badge_class = "badge-green" if tipo_label == "P.E" else "badge-blue"
                        nome_fan = cadastro.get('nome_fantasia', 'Sem nome')
                        num_cad = cadastro.get('numero_cadastro', 'N/A')
                        endereco_cad = cadastro.get('endereco', 'N/A')
                        lat_cad = cadastro.get('latitude', 'N/A')
                        lon_cad = cadastro.get('longitude', 'N/A')
                        quart_cad = cadastro.get('quarteirao', 'N/A')

                        st.markdown(f"""
                            <div class="sys-card">
                                <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 14px;">
                                    <div style="display: flex; align-items: center; gap: 10px;">
                                        <span class="sys-badge {badge_class}">{tipo_label}</span>
                                        <span class="sys-badge badge-orange">{freq_label}</span>
                                        <span style="font-size: 1.1rem; font-weight: 600; color: #2C3E50;">{nome_fan}</span>
                                    </div>
                                    <span style="font-size: 0.85rem; color: #85929E;">No {num_cad}</span>
                                </div>
                                <div class="info-grid">
                                    <div class="info-item">
                                        <div class="info-label">Endereco</div>
                                        <div class="info-value">{endereco_cad}</div>
                                    </div>
                                    <div class="info-item">
                                        <div class="info-label">Quarteirao</div>
                                        <div class="info-value">{quart_cad}</div>
                                    </div>
                                    <div class="info-item">
                                        <div class="info-label">Latitude</div>
                                        <div class="info-value">{lat_cad}</div>
                                    </div>
                                    <div class="info-item">
                                        <div class="info-label">Longitude</div>
                                        <div class="info-value">{lon_cad}</div>
                                    </div>
                                </div>
                            </div>
                        """, unsafe_allow_html=True)

                        if st.button(f"🗑️ Deletar  —  {nome_fan}", key=f"del_pe_ie_{idx}"):
                            db.reference(f'pe_ie_cadastros/{idx}').delete()
                            log_atividade(st.session_state.get('username'), f"Deletou {tipo_label}", f"No Cadastro: {num_cad}, Nome: {nome_fan}")
                            st.success(f"Cadastro '{nome_fan}' deletado com sucesso.")
                            st.cache_data.clear()
                            st.rerun()

            else:
                st.markdown("""
                    <div class="empty-state">
                        <div class="icon">🏗️</div>
                        <p>Nenhum P.E ou I.E cadastrado ainda.<br>Use a aba <strong>"Cadastrar Imovel"</strong> para comecar.</p>
                    </div>
                """, unsafe_allow_html=True)

        # =============================================
        # SUB-ABA 4: HISTORICO DE BOLETINS P.E / I.E
        # =============================================
        with sub_tab_historico:

            if not df_boletins_pe_ie.empty:
                df_boletins_pe_ie['data_dt'] = pd.to_datetime(df_boletins_pe_ie['data']).dt.date
                df_boletins_pe_ie_sorted = df_boletins_pe_ie.sort_values(by='data_dt', ascending=False)

                for idx_bol, boletim in df_boletins_pe_ie_sorted.iterrows():
                    data_fmt = pd.to_datetime(boletim['data']).strftime('%d/%m/%Y')
                    qtd_imoveis = len(boletim.get('imoveis_trabalhados', []))
                    criado_por = boletim.get('criado_por', 'N/A')

                    imoveis_list = boletim.get('imoveis_trabalhados', [])
                    imoveis_chips = ""
                    if imoveis_list and isinstance(imoveis_list, list):
                        for im in imoveis_list:
                            imoveis_chips += f'<span class="sys-chip">{im}</span>'

                    equipes_bol = boletim.get('equipes', [])
                    equipes_html = ""
                    if equipes_bol and isinstance(equipes_bol, list):
                        for eq_i, equipe in enumerate(equipes_bol):
                            if isinstance(equipe, dict):
                                membros_fmt = [formatar_nome(m) for m in equipe.get('membros', [])]
                                membros_tags = "".join([f'<span class="sys-tag">{m}</span>' for m in membros_fmt])
                                equipes_html += f'<div style="margin-top: 8px;"><span style="font-weight:600; color:#1B4F72; font-size:0.85rem;">Equipe {eq_i+1}:</span> {membros_tags}</div>'

                    obs = boletim.get('observacoes', '')
                    obs_html = f'<div style="margin-top:12px; padding:10px 14px; background:#FEF9E7; border-radius:8px; font-size:0.88rem; color:#7D6608;"><strong>Obs:</strong> {obs}</div>' if obs else ""

                    st.markdown(f"""
                        <div class="hist-card">
                            <div style="display: flex; justify-content: space-between; align-items: center;">
                                <div>
                                    <div class="info-value" style="font-size:1.05rem;font-weight:600;color:#1B4F72">📅 {data_fmt}</div>
                                    <div class="info-label">{qtd_imoveis} imovel(is) &middot; Criado por: {criado_por}</div>
                                </div>
                            </div>
                            <div class="sys-divider"></div>
                            <div style="margin-bottom:8px;">
                                <span style="font-weight:600; font-size:0.85rem; color:#566573;">IMOVEIS TRABALHADOS</span>
                            </div>
                            <div>{imoveis_chips if imoveis_chips else '<span style="color:#85929E;">Nenhum</span>'}</div>
                            {equipes_html}
                            {obs_html}
                        </div>
                    """, unsafe_allow_html=True)

                    if st.button(f"🗑️ Deletar boletim de {data_fmt}", key=f"del_bol_pe_{idx_bol}"):
                        db.reference(f'boletins_pe_ie/{idx_bol}').delete()
                        log_atividade(st.session_state.get('username'), "Deletou boletim P.E/I.E", f"Data: {data_fmt}")
                        st.success("Boletim deletado com sucesso.")
                        st.cache_data.clear()
                        st.rerun()

            else:
                st.markdown("""
                    <div class="empty-state">
                        <div class="icon">📭</div>
                        <p>Nenhum boletim de P.E/I.E registrado ainda.<br>Use a aba <strong>"Criar Boletim"</strong> para comecar.</p>
                    </div>
                """, unsafe_allow_html=True)

    with tab2:
        st.subheader("Visualizar e Editar Boletim")
        if df_boletins.empty:
            st.info("Nenhum boletim encontrado para visualização.")
        else:
            lista_boletins = sorted(df_boletins['id'].tolist(), reverse=True)
            boletim_id_selecionado = st.selectbox(
                "Selecione a data do boletim", 
                options=lista_boletins,
                format_func=lambda x: pd.to_datetime(x).strftime('%d/%m/%Y')
            )

            if boletim_id_selecionado:
                dados_boletim = df_boletins.loc[boletim_id_selecionado]
                st.markdown(f"#### Detalhes do dia: {pd.to_datetime(dados_boletim['data']).strftime('%d/%m/%Y')}")

                st.markdown(f"**Bairros trabalhados:** {dados_boletim.get('bairros', 'Não informado')}")
                st.markdown(f"**Atividades gerais:** {', '.join(dados_boletim.get('atividades_gerais', []))}")
                st.markdown(f"**Motorista(s):** {', '.join(map(formatar_nome, dados_boletim.get('motoristas', [])))}")
                st.markdown(f"**Ausentes (Manhã):** {', '.join(map(formatar_nome, dados_boletim.get('faltas_manha', {}).get('nomes', [])))} - *Motivo: {dados_boletim.get('faltas_manha', {}).get('motivo', '')}*")
                st.markdown(f"**Ausentes (Tarde):** {', '.join(map(formatar_nome, dados_boletim.get('faltas_tarde', {}).get('nomes', [])))} - *Motivo: {dados_boletim.get('faltas_tarde', {}).get('motivo', '')}*")
                
                st.divider()
                
                report_bytes = create_boletim_word_report(dados_boletim)
                st.download_button(
                    label="📥 Baixar Boletim (.docx)",
                    data=report_bytes,
                    file_name=f"Boletim_Diario_{boletim_id_selecionado}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                st.divider()

                col_manha, col_tarde = st.columns(2)
                with col_manha:
                    st.markdown("**Equipes da Manhã**")
                    equipes_manha = dados_boletim.get('equipes_manha', [])
                    if equipes_manha and isinstance(equipes_manha, list):
                        for i, equipe in enumerate(equipes_manha):
                            with st.expander(f"Equipe {i+1} (Manhã)"):
                                st.write(f"**Membros:** {', '.join(map(formatar_nome, equipe.get('membros', [])))}")
                                st.write(f"**Atividades:** {', '.join(equipe.get('atividades', []))}")
                                st.write(f"**Quarteirões:** {', '.join(equipe.get('quarteiroes', []))}")
                    else:
                        st.write("Nenhuma equipe registrada para a manhã.")

                with col_tarde:
                    st.markdown("**Equipes da Tarde**")
                    equipes_tarde = dados_boletim.get('equipes_tarde', [])
                    if equipes_tarde and isinstance(equipes_tarde, list):
                        for i, equipe in enumerate(equipes_tarde):
                            with st.expander(f"Equipe {i+1} (Tarde)"):
                                st.write(f"**Membros:** {', '.join(map(formatar_nome, equipe.get('membros', [])))}")
                                st.write(f"**Atividades:** {', '.join(equipe.get('atividades', []))}")
                                st.write(f"**Quarteirões:** {', '.join(equipe.get('quarteiroes', []))}")
                    else:
                        st.write("Nenhuma equipe registrada para a tarde.")
                
                st.divider()

                with st.expander("✏️ Editar este Boletim"):
                    with st.form(key="edit_boletim_form"):
                        st.warning("A edição de equipes ainda não é suportada. Em breve!")

                        bairros_edit = st.text_area("Bairros a serem trabalhados", value=dados_boletim.get('bairros', ''))
                        atividades_gerais_edit = st.multiselect("Atividades Gerais do Dia", options=["Controle de criadouros", "Visita a Imóveis", "ADL", "Nebulização"], default=dados_boletim.get('atividades_gerais', []))
                        
                        nome_map_full = {formatar_nome(nome): nome for nome in df_funcionarios['nome']}
                        lista_nomes_curtos_full_edit = sorted(list(nome_map_full.keys()))

                        motoristas_edit_curtos = st.multiselect("Motorista(s)", options=lista_nomes_curtos_full_edit, default=[formatar_nome(nome) for nome in dados_boletim.get('motoristas', [])])
                        
                        st.markdown("**Editar Faltas**")
                        faltas_manha_edit_curtos = st.multiselect("Ausentes (Manhã)", options=lista_nomes_curtos_full_edit, default=[formatar_nome(nome) for nome in dados_boletim.get('faltas_manha', {}).get('nomes', [])])
                        motivo_manha_edit = st.text_input("Motivo (Manhã)", value=dados_boletim.get('faltas_manha', {}).get('motivo', ''))
                        faltas_tarde_edit_curtos = st.multiselect("Ausentes (Tarde)", options=lista_nomes_curtos_full_edit, default=[formatar_nome(nome) for nome in dados_boletim.get('faltas_tarde', {}).get('nomes', [])])
                        motivo_tarde_edit = st.text_input("Motivo (Tarde)", value=dados_boletim.get('faltas_tarde', {}).get('motivo', ''))

                        submit_button = st.form_submit_button(label='Salvar Alterações')

                        if submit_button:
                            motoristas_completos_edit = [nome_map_full[nome] for nome in motoristas_edit_curtos]
                            faltas_manha_completos_edit = [nome_map_full[nome] for nome in faltas_manha_edit_curtos]
                            faltas_tarde_completos_edit = [nome_map_full[nome] for nome in faltas_tarde_edit_curtos]

                            dados_atualizados = {
                                "bairros": bairros_edit,
                                "atividades_gerais": atividades_gerais_edit,
                                "motoristas": motoristas_completos_edit,
                                "faltas_manha": {"nomes": faltas_manha_completos_edit, "motivo": motivo_manha_edit},
                                "faltas_tarde": {"nomes": faltas_tarde_completos_edit, "motivo": motivo_tarde_edit},
                            }
                            try:
                                ref = db.reference(f'boletins/{boletim_id_selecionado}')
                                ref.update(dados_atualizados)

                                log_atividade(st.session_state.get('username'), "Editou boletim", f"Boletim: {boletim_id_selecionado}")

                                st.success("Boletim atualizado com sucesso!")
                                st.cache_data.clear()
                                st.rerun()
                            except Exception as e:
                                st.error(f"Erro ao atualizar o boletim: {e}")

    with tab3:
        st.subheader("Mapa de Atividades por Dia")
        if df_boletins.empty or df_geo_quarteiroes.empty:
            st.warning("Dados de boletins ou geolocalização de quarteirões não estão disponíveis.")
        else:
            data_mapa = st.date_input("Selecione a data para visualizar no mapa", date.today(), key="mapa_data_plotly")
            boletim_id_mapa = data_mapa.strftime("%Y-%m-%d")

            if boletim_id_mapa in df_boletins.index:
                dados_boletim_mapa = df_boletins.loc[boletim_id_mapa]
                
                atividades_locs = []
                for turno in ['equipes_manha', 'equipes_tarde']:
                    if turno in dados_boletim_mapa and dados_boletim_mapa[turno] and isinstance(dados_boletim_mapa[turno], list):
                        for equipe in dados_boletim_mapa[turno]:
                            membros_nomes = ", ".join(map(formatar_nome, equipe.get('membros', [])))
                            for quarteirao in equipe.get('quarteiroes', []):
                                for atividade in equipe.get('atividades', ["Não especificada"]):
                                    atividades_locs.append({
                                        "quarteirao": quarteirao,
                                        "atividade": atividade,
                                        "equipe": membros_nomes
                                    })
                
                if not atividades_locs:
                    st.info(f"Nenhuma atividade de campo registrada para o dia {data_mapa.strftime('%d/%m/%Y')}.")
                else:
                    df_atividades_mapa = pd.DataFrame(atividades_locs)
                    df_mapa_final = pd.merge(df_atividades_mapa, df_geo_quarteiroes, left_on='quarteirao', right_on='quadra', how='inner')

                    if df_mapa_final.empty:
                        st.warning("Não foi possível encontrar as coordenadas para os quarteirões trabalhados.")
                    else:
                        st.info(f"Exibindo {len(df_mapa_final)} atividades no mapa para {data_mapa.strftime('%d/%m/%Y')}.")

                        fig = px.scatter_mapbox(
                            df_mapa_final,
                            lat="lat",
                            lon="lon",
                            color="atividade",
                            hover_name="quarteirao",
                            hover_data={"equipe": True, "atividade": True, "lat": False, "lon": False},
                            zoom=13,
                            height=600,
                            title="Distribuição de Atividades por Quarteirão"
                        )
                        
                        fig.update_layout(
                            mapbox_style="open-street-map",
                            margin={"r":0, "t":40, "l":0, "b":0},
                            legend_title_text='Atividades'
                        )
                        
                        st.plotly_chart(fig, use_container_width=True)
            else:
                st.info(f"Nenhum boletim encontrado para o dia {data_mapa.strftime('%d/%m/%Y')}.")

    with tab4:
        st.subheader("Dashboard de Produtividade")
        if df_boletins.empty:
            st.info("Nenhum dado de boletim para gerar o dashboard.")
        else:
            hoje = date.today()
            inicio_padrao = hoje - timedelta(days=30)
            data_inicio_dash, data_fim_dash = st.date_input(
                "Selecione o período de análise",
                [inicio_padrao, hoje],
                max_value=hoje,
                key="dash_date_range"
            )

            if data_inicio_dash and data_fim_dash and data_inicio_dash <= data_fim_dash:
                df_boletins['data_dt'] = pd.to_datetime(df_boletins['data']).dt.date
                df_boletins_filtrado = df_boletins[
                    (df_boletins['data_dt'] >= data_inicio_dash) &
                    (df_boletins['data_dt'] <= data_fim_dash)
                ]

                if df_boletins_filtrado.empty:
                    st.warning("Nenhum boletim encontrado no período selecionado.")
                else:
                    dados_analise = []
                    for _, boletim in df_boletins_filtrado.iterrows():
                        data = boletim['data']
                        for turno in ['equipes_manha', 'equipes_tarde']:
                            if turno in boletim and boletim[turno] and isinstance(boletim[turno], list):
                                for equipe in boletim[turno]:
                                    membros = equipe.get('membros', [])
                                    atividades = equipe.get('atividades', [])
                                    quarteiroes = equipe.get('quarteiroes', [])
                                    for membro in membros:
                                        dados_analise.append({'data': data, 'membro': formatar_nome(membro)})
                                    for atividade in atividades:
                                        dados_analise.append({'data': data, 'atividade': atividade})
                                    for quarteirao in quarteiroes:
                                        dados_analise.append({'data': data, 'quarteirao': quarteirao})

                    if not dados_analise:
                        st.info("Nenhuma atividade registrada no período para análise.")
                    else:
                        df_analise = pd.DataFrame(dados_analise)

                        st.divider()
                        st.markdown("#### Análise Gráfica do Período")
                        col1, col2 = st.columns(2)

                        with col1:
                            st.markdown("**Quarteirões Mais Trabalhados**")
                            if 'quarteirao' in df_analise.columns and not df_analise['quarteirao'].dropna().empty:
                                top_quarteiroes = df_analise['quarteirao'].value_counts().nlargest(15)
                                fig = px.bar(top_quarteiroes, x=top_quarteiroes.index, y=top_quarteiroes.values,
                                             labels={'y': 'Nº de Vezes Trabalhado', 'x': 'Quarteirão'}, text_auto=True)
                                fig.update_layout(title_x=0.5, xaxis_title="", yaxis_title="")
                                st.plotly_chart(fig, use_container_width=True)
                            else:
                                st.info("Nenhum dado de quarteirão no período.")
                        
                        with col2:
                            st.markdown("**Atividades Mais Executadas**")
                            if 'atividade' in df_analise.columns and not df_analise['atividade'].dropna().empty:
                                top_atividades = df_analise['atividade'].value_counts()
                                fig_pie = px.pie(top_atividades, values=top_atividades.values, names=top_atividades.index, 
                                                 hole=.3, color_discrete_sequence=px.colors.sequential.RdBu)
                                fig_pie.update_layout(title_x=0.5)
                                st.plotly_chart(fig_pie, use_container_width=True)
                            else:
                                st.info("Nenhum dado de atividade no período.")
                        
                        st.divider()
                        st.markdown("**Participação dos Funcionários (por turnos trabalhados)**")
                        if 'membro' in df_analise.columns and not df_analise['membro'].dropna().empty:
                            participacao = df_analise['membro'].value_counts()
                            fig_part = px.bar(participacao, x=participacao.index, y=participacao.values,
                                             labels={'y': 'Nº de Turnos', 'x': 'Funcionário'}, text_auto=True)
                            fig_part.update_layout(title_x=0.5, xaxis_title="", yaxis_title="")
                            st.plotly_chart(fig_part, use_container_width=True)
                        else:
                            st.info("Nenhum dado de participação no período.")

# ### NOVO MÓDULO DE LOGS ###
def modulo_logs():
    """Renderiza a pagina para visualizar os logs de atividade."""
    st.markdown("""
        <div class="mod-header">
            <h2>📄 Logs de Atividade</h2>
            <p>Historico completo de acoes realizadas no sistema</p>
        </div>
    """, unsafe_allow_html=True)

    df_logs = carregar_dados_firebase('logs_de_atividade')

    if not df_logs.empty:
        df_logs['timestamp_dt'] = pd.to_datetime(df_logs['timestamp'])
        df_logs_display = df_logs.sort_values(by='timestamp_dt', ascending=False).copy()
        
        # Opcional: Filtros para facilitar a busca
        st.sidebar.markdown("---")
        st.sidebar.subheader("Filtrar Logs")
        usuarios_logados = sorted(df_logs_display['usuario'].unique().tolist())
        filtro_usuario = st.sidebar.selectbox("Filtrar por Usuário", options=["Todos"] + usuarios_logados)

        acoes_disponiveis = sorted(df_logs_display['acao'].unique().tolist())
        filtro_acao = st.sidebar.multiselect("Filtrar por Ação", options=acoes_disponiveis)

        if filtro_usuario != "Todos":
            df_logs_display = df_logs_display[df_logs_display['usuario'] == filtro_usuario]

        if filtro_acao:
            df_logs_display = df_logs_display[df_logs_display['acao'].isin(filtro_acao)]

        # Exibe os dados filtrados
        cols_to_display = ['timestamp', 'usuario', 'acao', 'detalhes']
        st.dataframe(
            df_logs_display[cols_to_display].rename(
                columns={'timestamp': 'Data/Hora', 'usuario': 'Usuário', 'acao': 'Ação', 'detalhes': 'Detalhes'}
            ),
            use_container_width=True,
            hide_index=True
        )
    else:
        st.info("Nenhum log de atividade encontrado.")


# --- ESTRUTURA PRINCIPAL DA APLICAÇÃO (LOGIN E NAVEGAÇÃO) ---

def login_screen():
    """Renderiza a tela de login."""
    st.markdown("""
        <div class="login-container">
            <div class="login-logo">🏥</div>
            <div class="login-title">Sistema de Gestao</div>
            <div class="login-subtitle">Vigilancia Epidemiologica — Guaratingueta/SP</div>
        </div>
    """, unsafe_allow_html=True)
    
    col_spacer1, col_login, col_spacer2 = st.columns([1.2, 1, 1.2])
    with col_login:
        with st.form("login_form"):
            st.markdown('<div class="sys-card">', unsafe_allow_html=True)
            username = st.text_input("Usuario", key="login_username")
            password = st.text_input("Senha", type="password", key="login_password")
            submit_button = st.form_submit_button("Entrar", use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)
            if submit_button:
                if username in USERS and USERS[username] == password:
                    st.session_state['logged_in'] = True
                    st.session_state['username'] = username
                    st.rerun()
                else:
                    st.error("Usuario ou senha invalidos.")

def main_app():
    """Controla a navegação e a exibição dos módulos após o login."""
    if 'evento_para_editar_id' not in st.session_state:
        st.session_state.evento_para_editar_id = None

    if st.session_state.get('module_choice'):
        with st.sidebar:
            st.title("Navegação")
            st.write(f"Usuário: **{st.session_state['username']}**")
            st.divider()
            if st.button("⬅️ Voltar ao Painel de Controle"):
                st.session_state.evento_para_editar_id = None
                st.session_state['module_choice'] = None
                st.rerun()
            st.divider()
            if st.button("Logout"):
                log_atividade(st.session_state.get('username'), "Logout", "Usuário encerrou a sessão.")
                for key in list(st.session_state.keys()):
                    del st.session_state[key]
                st.rerun()

        if st.session_state['module_choice'] == "Denúncias":
            modulo_denuncias()
        elif st.session_state['module_choice'] == "Recursos Humanos":
            modulo_rh()
        elif st.session_state['module_choice'] == "Boletim":
            modulo_boletim()
        elif st.session_state['module_choice'] == "Logs":
            modulo_logs()

    else:
        st.markdown("""
            <div class="mod-header">
                <h2>🏥 Painel de Controle</h2>
                <p>Bem-vindo(a), """ + st.session_state['username'] + """! Selecione um modulo abaixo para comecar.</p>
            </div>
        """, unsafe_allow_html=True)
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.markdown("""
                <div class="module-card">
                    <div class="icon">🚨</div>
                    <div class="title">Denuncias</div>
                    <div class="desc">Registrar e gerenciar</div>
                </div>
            """, unsafe_allow_html=True)
            if st.button("Acessar Denuncias", use_container_width=True, key="btn_den"):
                st.session_state['module_choice'] = "Denúncias"
                st.rerun()
        with col2:
            st.markdown("""
                <div class="module-card">
                    <div class="icon">👥</div>
                    <div class="title">Recursos Humanos</div>
                    <div class="desc">Equipe e ferias</div>
                </div>
            """, unsafe_allow_html=True)
            if st.button("Acessar RH", use_container_width=True, key="btn_rh"):
                st.session_state['module_choice'] = "Recursos Humanos"
                st.rerun()
        with col3:
            st.markdown("""
                <div class="module-card">
                    <div class="icon">🗓️</div>
                    <div class="title">Boletim Diario</div>
                    <div class="desc">Programacao e equipes</div>
                </div>
            """, unsafe_allow_html=True)
            if st.button("Acessar Boletim", use_container_width=True, key="btn_bol"):
                st.session_state['module_choice'] = "Boletim"
                st.rerun()
        with col4:
            st.markdown("""
                <div class="module-card">
                    <div class="icon">📄</div>
                    <div class="title">Logs de Atividade</div>
                    <div class="desc">Historico de acoes</div>
                </div>
            """, unsafe_allow_html=True)
            if st.button("Acessar Logs", use_container_width=True, key="btn_log"):
                st.session_state['module_choice'] = "Logs"
                st.rerun()
        
        st.markdown('<div class="sys-divider"></div>', unsafe_allow_html=True)

        col_form, col_cal = st.columns([1, 1.5])

        with col_form:
            df_funcionarios = carregar_dados_firebase('funcionarios')
            if not df_funcionarios.empty:
                lista_nomes_curtos = sorted([formatar_nome(nome) for nome in df_funcionarios['nome']])
            else:
                lista_nomes_curtos = []

            if st.session_state.evento_para_editar_id:
                st.subheader("✏️ Editando Evento")
                df_avisos_all = carregar_dados_firebase('avisos')
                dados_evento = df_avisos_all.loc[st.session_state.evento_para_editar_id]

                with st.form("form_avisos_edit"):
                    titulo_edit = st.text_input("Título do Evento", value=dados_evento.get('titulo', ''))
                    
                    tipos_de_evento = ["Aviso", "Compromisso", "Reunião", "Curso", "Educativa"]
                    tipo_idx = tipos_de_evento.index(dados_evento.get('tipo_aviso')) if dados_evento.get('tipo_aviso') in tipos_de_evento else 0
                    tipo_edit = st.selectbox("Tipo", tipos_de_evento, index=tipo_idx, key='tipo_evento_edit')

                    data_val = pd.to_datetime(dados_evento.get('data')).date() if pd.notna(dados_evento.get('data')) else date.today()
                    data_edit = st.date_input("Data", value=data_val)

                    participantes_edit = []
                    if tipo_edit in ["Reunião", "Curso", "Educativa"]:
                        participantes_edit = st.multiselect("Participantes", options=lista_nomes_curtos, default=dados_evento.get('participantes', []))
                    
                    descricao_edit = st.text_area("Descrição (Opcional)", value=dados_evento.get('descricao', ''))

                    col_save, col_cancel = st.columns(2)
                    with col_save:
                        if st.form_submit_button("Salvar Alterações", use_container_width=True):
                            dados_atualizados = {
                                'titulo': titulo_edit,
                                'data': data_edit.strftime("%Y-%m-%d"),
                                'tipo_aviso': tipo_edit,
                                'descricao': descricao_edit,
                                'participantes': participantes_edit
                            }
                            db.reference(f'avisos/{st.session_state.evento_para_editar_id}').update(dados_atualizados)
                            
                            log_atividade(st.session_state.get('username'), "Editou aviso no mural", f"Título: {titulo_edit}")

                            st.success("Evento atualizado com sucesso!")
                            st.session_state.evento_para_editar_id = None
                            st.cache_data.clear()
                            st.rerun()

                    with col_cancel:
                        if st.form_submit_button("Cancelar", type="secondary", use_container_width=True):
                            st.session_state.evento_para_editar_id = None
                            st.rerun()
                st.divider()


            st.subheader("📝 Adicionar no Mural")
            
            tipos_de_evento_add = ["Aviso", "Compromisso", "Reunião", "Curso", "Educativa"]
            aviso_tipo_add = st.selectbox("Tipo de Evento", tipos_de_evento_add, key='tipo_evento_selecionado')
            
            with st.form("form_avisos_add", clear_on_submit=True):
                aviso_titulo = st.text_input("Título do Evento")
                aviso_data = st.date_input("Data")
                
                participantes = []
                if st.session_state.tipo_evento_selecionado in ["Reunião", "Curso", "Educativa"]:
                    participantes = st.multiselect("Participantes", options=lista_nomes_curtos)

                aviso_descricao = st.text_area("Descrição (Opcional)")
                
                submitted = st.form_submit_button("Salvar no Mural")
                if submitted:
                    if aviso_titulo and aviso_data:
                        try:
                            aviso_id = str(int(time.time() * 1000))
                            ref = db.reference(f'avisos/{aviso_id}')
                            ref.set({
                                'titulo': aviso_titulo,
                                'data': aviso_data.strftime("%Y-%m-%d"),
                                'tipo_aviso': st.session_state.tipo_evento_selecionado,
                                'descricao': aviso_descricao,
                                'participantes': participantes 
                            })
                            
                            log_atividade(st.session_state.get('username'), "Adicionou aviso no mural", f"Título: {aviso_titulo}")

                            st.success("Evento salvo no mural com sucesso!")
                            st.cache_data.clear()
                            st.rerun()
                        except Exception as e:
                            st.error(f"Erro ao salvar o aviso: {e}")
                    else:
                        st.warning("Por favor, preencha o Título e a Data.")

            st.divider()
            st.subheader("🗓️ Eventos Agendados")
            df_avisos = carregar_dados_firebase('avisos')

            filtro_data = st.date_input("Filtrar eventos por dia", value=date.today())

            if not df_avisos.empty:
                df_avisos['data_dt'] = pd.to_datetime(df_avisos['data']).dt.date
                avisos_filtrados = df_avisos[df_avisos['data_dt'] == filtro_data].sort_values(by='titulo')

                if avisos_filtrados.empty:
                    st.info(f"Nenhum evento agendado para {filtro_data.strftime('%d/%m/%Y')}.")
                else:
                    for id, aviso in avisos_filtrados.iterrows():
                        with st.expander(f"{aviso.get('tipo_aviso', 'Evento')}: **{aviso.get('titulo', 'Sem título')}**"):
                            if aviso.get('descricao'):
                                st.markdown(f"**Descrição:** {aviso.get('descricao')}")
                            
                            lista_participantes = aviso.get('participantes', [])
                            if lista_participantes and isinstance(lista_participantes, list):
                                participantes_validos = [str(p) for p in lista_participantes if p]
                                if participantes_validos:
                                    st.markdown(f"**Participantes:** {', '.join(participantes_validos)}")
                            
                            st.markdown("---")
                            col_b1, col_b2, _ = st.columns([1, 1, 3])
                            with col_b1:
                                if st.button("✏️ Editar", key=f"edit_{id}", use_container_width=True):
                                    st.session_state.evento_para_editar_id = id
                                    st.rerun()
                            with col_b2:
                                if st.button("🗑️ Deletar", key=f"del_{id}", type="primary", use_container_width=True):
                                    db.reference(f'avisos/{id}').delete()
                                    
                                    log_atividade(st.session_state.get('username'), "Deletou aviso no mural", f"Título: {aviso.get('titulo')}")

                                    st.success(f"Evento '{aviso.get('titulo')}' deletado.")
                                    st.cache_data.clear()
                                    st.rerun()
            else:
                st.info("Nenhum evento no mural para exibir.")


        with col_cal:
            st.subheader("📅 Calendário Geral de Eventos e Ausências")
            
            df_folgas = carregar_dados_firebase('folgas_ferias')
            df_avisos_cal = carregar_dados_firebase('avisos')
            
            calendar_events = []

            if not df_folgas.empty:
                for _, row in df_folgas.iterrows():
                    calendar_events.append({
                        "title": f"AUSÊNCIA: {formatar_nome(row['nome_funcionario'])} ({row['tipo']})",
                        "start": row['data_inicio'],
                        "end": (pd.to_datetime(row['data_fim']) + timedelta(days=1)).strftime("%Y-%m-%d"),
                        "color": "#FF4B4B" if row['tipo'] == "Férias" else "#FFA07A",
                    })
            
            if not df_avisos_cal.empty:
                event_colors = {
                    "Aviso": "#ffc107",
                    "Compromisso": "#28a745",
                    "Reunião": "#007bff",
                    "Curso": "#6f42c1",
                    "Educativa": "#fd7e14"
                }
                for _, row in df_avisos_cal.iterrows():
                    event_type = row.get('tipo_aviso', 'Aviso')
                    calendar_events.append({
                        "title": f"{event_type.upper()}: {row['titulo']}",
                        "start": row['data'],
                        "end": (pd.to_datetime(row['data']) + timedelta(days=1)).strftime("%Y-%m-%d"),
                        "color": event_colors.get(event_type, "#6c757d"),
                    })

            calendar_options = {
                "initialView": "dayGridMonth",
                "height": "800px",
                "locale": "pt-br",
                "headerToolbar": {
                    "left": "prev,next today",
                    "center": "title",
                    "right": "dayGridMonth,timeGridWeek"
                },
                "eventTimeFormat": {
                    "hour": '2-digit',
                    "minute": '2-digit',
                    "meridiem": False
                }
            }
            
            if calendar_events:
                calendar(events=calendar_events, options=calendar_options, key="calendario_mural_atualizado")
            else:
                st.info("Nenhum evento no mural ou ausência registrada para exibir no calendário.")

if __name__ == "__main__":
    if 'logged_in' not in st.session_state:
        st.session_state['logged_in'] = False
    if st.session_state['logged_in']:
        main_app()
    else:
        login_screen()
